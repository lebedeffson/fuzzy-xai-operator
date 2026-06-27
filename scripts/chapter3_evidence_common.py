from __future__ import annotations

import csv
import json
import math
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd
import shap
import yaml
from docx import Document
from lime.lime_tabular import LimeTabularExplainer
from sklearn.datasets import load_breast_cancer
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import accuracy_score
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import StandardScaler


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "reports/chapter3"
PATCH = ROOT / "patches"
CONFIG_DIR = ROOT / "configs/chapter3"
DOCX_IN = Path("/mnt/data/03_glava_3_final_reviewer_fixed_submission_ready.docx")
LOCAL_DOCX = ROOT / "docs/chapter3/03_glava_3_final_reviewer_fixed_submission_ready.docx"
CONFIG_PATH = CONFIG_DIR / "real_conflict_experiment.yaml"

BAD_FH_TERMS = [
    "F_H",
    "FH",
    "hesitant",
    "Hesitant",
    "хезитант",
    "хезитантное",
    "множество альтернатив",
    "альтернативные экспертные оценки",
]

ACTIONS = ["accept", "lower_confidence", "request_more_data", "defer_to_human", "block"]


def ensure_dirs() -> None:
    REPORT.mkdir(parents=True, exist_ok=True)
    PATCH.mkdir(parents=True, exist_ok=True)
    CONFIG_DIR.mkdir(parents=True, exist_ok=True)


def read_config() -> dict[str, Any]:
    data = yaml.safe_load(CONFIG_PATH.read_text(encoding="utf-8"))
    return dict(data)


def dump_yaml(data: Any, indent: int = 0) -> str:
    pad = " " * indent
    if isinstance(data, dict):
        lines = []
        for k, v in data.items():
            if isinstance(v, (dict, list)):
                lines.append(f"{pad}{k}:")
                lines.append(dump_yaml(v, indent + 2).rstrip())
            else:
                lines.append(f"{pad}{k}: {v}")
        return "\n".join(lines) + "\n"
    if isinstance(data, list):
        return "\n".join(f"{pad}- {v}" for v in data) + "\n"
    return f"{pad}{data}\n"


def write_csv(path: Path, rows: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fields = sorted({k for r in rows for k in r})
    with path.open("w", encoding="utf-8", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fields)
        w.writeheader()
        w.writerows(rows)


def load_docx_text(path: Path) -> list[str]:
    doc = Document(str(path))
    out = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                out.extend(p.text for p in cell.paragraphs if p.text.strip())
    return out


def audit_docx() -> None:
    ensure_dirs()
    src = LOCAL_DOCX if LOCAL_DOCX.exists() else DOCX_IN
    rows: list[dict[str, Any]] = []
    paragraphs: list[str] = []
    if src.exists():
        paragraphs = load_docx_text(src)
        for i, text in enumerate(paragraphs):
            for term in BAD_FH_TERMS:
                if term.lower() in text.lower():
                    section = next((m.group(0) for m in [re.match(r"\s*3(?:\.\d+)*", text)] if m), "")
                    dep = classify_dependency(text)
                    rows.append(
                        {
                            "section": section or "unknown",
                            "page_estimate": max(1, i // 5 + 1),
                            "paragraph_index": i,
                            "matched_text": term,
                            "context_before": paragraphs[i - 1][:240] if i else "",
                            "context_after": paragraphs[i + 1][:240] if i + 1 < len(paragraphs) else "",
                            "dependency_type": dep,
                            "recommended_action": recommended_action(dep, text),
                        }
                    )
    else:
        rows.append(
            {
                "section": "source_missing",
                "page_estimate": "",
                "paragraph_index": "",
                "matched_text": "DOCX not found",
                "context_before": "",
                "context_after": "",
                "dependency_type": "reference_only",
                "recommended_action": "move_to_appendix",
            }
        )
    write_csv(REPORT / "hesitant_dependency_map.csv", rows)
    (REPORT / "current_chapter_audit.md").write_text(
        "# Аудит текущей главы 3\n\n"
        f"- Исходный DOCX: `{src}`\n"
        f"- Статус: {'найден' if src.exists() else 'не найден'}\n"
        f"- Найдено зависимостей от `F_H`/hesitant: {0 if rows and rows[0]['section']=='source_missing' else len(rows)}\n"
        "- `F_H` не должен оставаться в минимальном ядре главы.\n",
        encoding="utf-8",
    )
    make_hesitant_plan(rows)


def classify_dependency(text: str) -> str:
    t = text.lower()
    if any(x in t for x in ["f_core", "ядр", "определ"]):
        return "definition"
    if any(x in t for x in ["=", "{", "формул", "rho", "delta"]):
        return "formula"
    if "таблиц" in t:
        return "table"
    if any(x in t for x in ["пример", "sample"]):
        return "example"
    if any(x in t for x in ["эксперимент", "результат"]):
        return "experiment"
    if "положение" in t:
        return "protected_position"
    return "reference_only"


def recommended_action(dep: str, text: str) -> str:
    if dep in {"definition", "formula"}:
        return "replace_by_NAS"
    if dep == "experiment":
        return "replace_by_F_ML"
    if dep == "table":
        return "replace_by_NAS"
    if dep in {"example", "protected_position"}:
        return "move_to_appendix"
    return "keep_as_optional_note"


def make_hesitant_plan(rows: list[dict[str, Any]]) -> None:
    n = 0 if rows and rows[0].get("section") == "source_missing" else len(rows)
    text = f"""# План удаления `F_H` из ядра главы 3

Найдено размеченных зависимостей: {n}.

## Новое минимальное ядро

`F_core = {{F0, F_int, NAS, F_ML}}`, где `NAS = F_N^src`.

Покрытие типов неопределённости:

| Класс | Покрытие |
| --- | --- |
| `F0` | `u_num`, `u_ling` |
| `F_int` | `u_int` |
| `NAS` | `u_non`, `u_conf`, `u_trace/source conflict` |
| `F_ML` | `u_time`, `u_shift`, `u_cf`, `u_user`, `u_multi`, combinations |

Если после удаления `F_H` остаётся `u_expert`, его следует включать в `NAS` как конфликт разных источников экспертных оценок либо выносить во внешний необязательный класс приложения.

## Действия по тексту

1. В разделах 3.3, 3.6, 3.7, 3.8, 3.9, 3.14, 3.15, 3.22 заменить минимальное ядро на `{{F0, F_int, NAS, F_ML}}`.
2. Формулы и таблицы, где `F_H` входит в `F_core`, переписать через `NAS`.
3. Примеры с hesitant-представлениями вынести в электронное приложение или заменить примером `SHAP/LIME/rule`-конфликта.
4. Оставить короткое примечание: Hesitant-представления могут рассматриваться как внешний расширяемый класс, но в минимальное ядро главы 3 не входят, поскольку не участвуют в финальной экспериментальной проверке.
5. В основном тексте не оставлять утверждения, что без `F_H` ядро не покрывает заявленное пространство.
"""
    (REPORT / "hesitant_removal_plan.md").write_text(text, encoding="utf-8")


@dataclass
class ModelBundle:
    x: np.ndarray
    y: np.ndarray
    names: list[str]
    scaler: StandardScaler
    model: LogisticRegression
    train_idx: np.ndarray
    val_idx: np.ndarray
    test_idx: np.ndarray
    x_train: np.ndarray


def train_model(seed: int = 42) -> ModelBundle:
    bc = load_breast_cancer()
    x = bc.data
    y = (bc.target == 0).astype(int)  # malignant = 1
    idx = np.arange(len(y))
    tr_idx, tmp_idx, y_tr, y_tmp = train_test_split(idx, y, test_size=0.40, stratify=y, random_state=seed)
    val_idx, test_idx, _, _ = train_test_split(tmp_idx, y_tmp, test_size=0.50, stratify=y_tmp, random_state=seed)
    scaler = StandardScaler().fit(x[tr_idx])
    model = LogisticRegression(max_iter=2000, solver="liblinear", random_state=seed)
    model.fit(scaler.transform(x[tr_idx]), y[tr_idx])
    return ModelBundle(x, y, list(bc.feature_names), scaler, model, tr_idx, val_idx, test_idx, scaler.transform(x[tr_idx]))


def predict_proba(bundle: ModelBundle, rows: np.ndarray) -> np.ndarray:
    return bundle.model.predict_proba(bundle.scaler.transform(rows))


def top_features(values: np.ndarray, names: list[str], k: int) -> tuple[list[str], dict[str, int]]:
    order = np.argsort(np.abs(values))[::-1][:k]
    feats = [names[i] for i in order]
    signs = {names[i]: int(np.sign(values[i])) for i in order}
    return feats, signs


def action_from_rho(rho: float, chi_crit: int = 0, thetas: tuple[float, float, float, float] = (0.25, 0.45, 0.65, 0.82)) -> str:
    if chi_crit:
        return "block"
    t1, t2, t3, t4 = thetas
    if rho < t1:
        return "accept"
    if rho < t2:
        return "lower_confidence"
    if rho < t3:
        return "request_more_data"
    if rho < t4:
        return "defer_to_human"
    return "block"


def build_real_conflicts() -> pd.DataFrame:
    ensure_dirs()
    cfg = read_config()
    seed, k = int(cfg["seed"]), int(cfg["top_k"])
    rng = np.random.default_rng(seed)
    bundle = train_model(seed)
    eval_pool = np.concatenate([bundle.val_idx, bundle.test_idx])
    chosen = rng.choice(eval_pool, size=int(cfg["n_evaluation_cases"]), replace=True)
    split_map = {int(i): "validation" for i in bundle.val_idx} | {int(i): "test" for i in bundle.test_idx}

    shap_explainer = shap.LinearExplainer(bundle.model, bundle.x_train)
    shap_values = shap_explainer(bundle.scaler.transform(bundle.x[eval_pool])).values
    shap_by_id = {int(i): shap_values[pos] for pos, i in enumerate(eval_pool)}

    lime_exp = LimeTabularExplainer(
        bundle.x[bundle.train_idx],
        feature_names=bundle.names,
        class_names=["benign", "malignant"],
        mode="classification",
        discretize_continuous=False,
        random_state=seed,
    )

    lime_by_id: dict[int, np.ndarray] = {}
    for row_id in sorted(set(int(i) for i in chosen)):
        exp = lime_exp.explain_instance(
            bundle.x[row_id],
            lambda z: predict_proba(bundle, z),
            labels=[1],
            num_features=len(bundle.names),
            num_samples=int(cfg["lime_num_samples"]),
        )
        vals = np.zeros(len(bundle.names))
        name_to_idx = {n: i for i, n in enumerate(bundle.names)}
        for feat, weight in exp.as_list(label=1):
            feat = str(feat)
            if feat in name_to_idx:
                vals[name_to_idx[feat]] = float(weight)
        lime_by_id[row_id] = vals

    rows: list[dict[str, Any]] = []
    auto_low, auto_high = float(cfg["auto_low"]), float(cfg["auto_high"])
    theta_rank = float(cfg["theta_rank"])
    for case_no, row_id in enumerate(chosen):
        row_id = int(row_id)
        xrow = bundle.x[[row_id]]
        proba = predict_proba(bundle, xrow)[0]
        p = float(proba[1])
        y_pred = int(p >= 0.5)
        shap_feats, shap_sign = top_features(shap_by_id[row_id], bundle.names, k)
        lime_feats, lime_sign = top_features(lime_by_id[row_id], bundle.names, k)
        common = set(shap_feats) & set(lime_feats)
        jacc = len(common) / len(set(shap_feats) | set(lime_feats))
        sign_disagree = sum(1 for f in common if shap_sign.get(f, 0) * lime_sign.get(f, 0) < 0)
        conflict_rank = int(jacc < theta_rank)
        conflict_sign = int(sign_disagree > 0)
        shap_support = float(np.mean([shap_sign[f] for f in shap_feats]))
        lime_support = float(np.mean([lime_sign[f] for f in lime_feats]))
        rule_action = "block" if p >= auto_high else ("accept" if p <= auto_low else "defer_to_human")
        local_high = shap_support + lime_support > 0.25
        local_low = shap_support + lime_support < -0.25
        conflict_rule = int((rule_action == "block" and local_low) or (rule_action == "accept" and local_high))
        chi_r = int(conflict_rank or conflict_sign or conflict_rule)
        chi_crit = int((conflict_sign and conflict_rule) or (chi_r and (p <= auto_low or p >= auto_high)))
        uncertainty = 1.0 - 2.0 * abs(p - 0.5)
        i_pre = max(0.0, min(1.0, jacc - 0.20 * sign_disagree))
        rho_f0 = p
        action_f0 = action_from_rho(rho_f0)
        i_nas = max(0.0, min(1.0, 1.0 - jacc + 0.20 * sign_disagree + 0.25 * conflict_rule))
        rho_nas = max(0.0, min(1.0, 0.55 * p + 0.25 * uncertainty + 0.20 * chi_r))
        action_nas = "block" if chi_crit else action_from_rho(rho_nas)
        rho_fml = max(0.0, min(1.0, 0.50 * rho_nas + 0.25 * i_nas + 0.25 * chi_r))
        action_fml = "block" if chi_crit else action_from_rho(rho_fml)
        delta = abs(rho_fml - rho_f0)
        final_action = action_fml
        rows.append(
            {
                "case_id": f"bc_boot_{case_no:04d}",
                "source_row_id": row_id,
                "split": split_map[row_id],
                "seed": seed,
                "y_true": int(bundle.y[row_id]),
                "y_pred": y_pred,
                "p_malignant": round(p, 6),
                "top_shap_features": "|".join(shap_feats),
                "top_lime_features": "|".join(lime_feats),
                "shap_signs": json.dumps(shap_sign, ensure_ascii=False),
                "lime_signs": json.dumps(lime_sign, ensure_ascii=False),
                "topk_jaccard": round(jacc, 6),
                "sign_disagreement_count": sign_disagree,
                "rule_action": rule_action,
                "f0_action": action_f0,
                "action_f0": action_f0,
                "nas_action": action_nas,
                "action_nas": action_nas,
                "fml_action": action_fml,
                "action_fml": action_fml,
                "conflict_rank": conflict_rank,
                "conflict_sign": conflict_sign,
                "conflict_rule": conflict_rule,
                "chi_R": chi_r,
                "chi_R_crit": chi_crit,
                "rho_f0": round(rho_f0, 6),
                "rho_nas": round(rho_nas, 6),
                "rho_fml": round(rho_fml, 6),
                "delta_reduction": round(delta, 6),
                "final_action": final_action,
                "T": round(max(0.0, p), 6),
                "I": round(i_nas, 6),
                "F": round(max(0.0, 1.0 - p), 6),
                "src_T": "model_probability+SHAP" if p >= 0.5 else "LIME/rule",
                "src_I": "SHAP/LIME/rule disagreement",
                "src_F": "counter-evidence source",
                "u_M": round(uncertainty, 6),
                "I_pre": round(i_pre, 6),
                "Delta_M": round(delta, 6),
            }
        )
    df = pd.DataFrame(rows)
    df.to_csv(REPORT / "real_conflict_summary.csv", index=False)
    write_real_conflict_report(df, bundle)
    return df


def write_real_conflict_report(df: pd.DataFrame, bundle: ModelBundle) -> None:
    acc = accuracy_score(bundle.y[bundle.test_idx], bundle.model.predict(bundle.scaler.transform(bundle.x[bundle.test_idx])))
    n_unique = int(df["source_row_id"].nunique())
    text = f"""# Эксперимент на реальных объяснительных конфликтах

Датасет: Breast Cancer Wisconsin. Это не клиническая апробация, а проверка XAI-конфликтов на реальной табличной задаче.

- Уникальных объектов в датасете: {len(bundle.y)}
- Уникальных объектов в evaluation cases: {n_unique}
- Evaluation cases: {len(df)}
- Формирование 1000 cases: bootstrap из validation/test split; это evaluation cases, а не уникальные клинические объекты.
- Модель: logistic regression, одна и та же модель для SHAP и LIME.
- Test accuracy модели: {acc:.4f}
- Реальный конфликт: rank/sign/rule расхождение между SHAP, LIME и rule/ExplainPlan-профилем.

| Метрика | Значение |
| --- | ---: |
| `n_cases` | {len(df)} |
| `n_unique_objects` | {n_unique} |
| `n_conflicts` | {int(df['chi_R'].sum())} |
| `n_critical_conflicts` | {int(df['chi_R_crit'].sum())} |
| `rank conflicts` | {int(df['conflict_rank'].sum())} |
| `sign conflicts` | {int(df['conflict_sign'].sum())} |
| `rule conflicts` | {int(df['conflict_rule'].sum())} |

Синтетические перевороты правил не используются.
"""
    (REPORT / "real_conflict_experiment_report.md").write_text(text, encoding="utf-8")


def load_conflicts() -> pd.DataFrame:
    path = REPORT / "real_conflict_summary.csv"
    if not path.exists():
        return build_real_conflicts()
    return pd.read_csv(path)


def binom_ci(success: int, n: int, z: float = 1.96) -> tuple[float, float]:
    if n <= 0:
        return (0.0, 0.0)
    p = success / n
    denom = 1 + z * z / n
    center = (p + z * z / (2 * n)) / denom
    half = z * math.sqrt((p * (1 - p) + z * z / (4 * n)) / n) / denom
    return max(0, center - half), min(1, center + half)


def f0_vs_nas() -> pd.DataFrame:
    ensure_dirs()
    df = load_conflicts()
    n = len(df)
    n_unique = int(df["source_row_id"].nunique())
    diff = df["action_f0"] != df["action_nas"]
    f0_accept_nas_block = (df["action_f0"] == "accept") & (df["action_nas"] == "block")
    ci_diff = binom_ci(int(diff.sum()), n)
    ci_ab = binom_ci(int(f0_accept_nas_block.sum()), n)
    row = {
        "n_cases": n,
        "n_unique_objects": n_unique,
        "n_conflicts": int(df["chi_R"].sum()),
        "n_critical_conflicts": int(df["chi_R_crit"].sum()),
        "n_action_diff_f0_nas": int(diff.sum()),
        "share_action_diff_f0_nas": diff.mean(),
        "ci95_action_diff_low": ci_diff[0],
        "ci95_action_diff_high": ci_diff[1],
        "n_f0_accept_nas_block": int(f0_accept_nas_block.sum()),
        "share_f0_accept_nas_block": f0_accept_nas_block.mean(),
        "ci95_f0_accept_nas_block_low": ci_ab[0],
        "ci95_f0_accept_nas_block_high": ci_ab[1],
    }
    out = pd.DataFrame([row])
    out.to_csv(REPORT / "f0_vs_nas_action_diff.csv", index=False)
    return out


def apply_config(df: pd.DataFrame, cfg: dict[str, Any]) -> pd.Series:
    w = cfg["weights"]
    th = cfg["thresholds"]
    rho = (
        w["w_p"] * df["p_malignant"]
        + w["w_u"] * df["u_M"]
        + w["w_I"] * (1.0 - df["I_pre"])
        + w["w_Delta"] * df["Delta_M"]
        + w["w_R"] * df["chi_R"]
    ).clip(0, 1)
    actions = []
    for r, crit in zip(rho, df["chi_R_crit"]):
        actions.append(action_from_rho(float(r), int(crit), tuple(th)))
    return pd.Series(actions, index=df.index)


def objective(df: pd.DataFrame, actions: pd.Series) -> dict[str, float]:
    missed = int(((df["chi_R_crit"] == 1) & (actions != "block")).sum())
    false_auto = int(((actions == "accept") & ((df["y_true"] != df["y_pred"]) | (df["chi_R"] == 1))).sum())
    unsafe = int(((actions == "accept") & (df["chi_R"] == 1)).sum())
    false_block = int(((actions == "block") & (df["y_true"] == df["y_pred"]) & (df["chi_R"] == 0) & (df["p_malignant"] < 0.35)).sum())
    excessive = int(((actions == "defer_to_human") & (df["chi_R"] == 0) & (df["p_malignant"] < 0.35)).sum())
    j = 5 * missed + 3 * false_auto + 2 * unsafe + false_block + 0.5 * excessive
    return {
        "missed_critical_ruptures": missed,
        "false_auto_accept": false_auto,
        "unsafe_accept_with_conflict": unsafe,
        "false_block": false_block,
        "excessive_defer": excessive,
        "objective_J": float(j),
    }


def split_eval(df: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    return df[df["split"] == "validation"].copy(), df[df["split"] == "test"].copy()


def normalize_weights(raw: np.ndarray) -> dict[str, float]:
    raw = np.maximum(raw, 0)
    raw = raw / raw.sum()
    keys = ["w_p", "w_u", "w_I", "w_Delta", "w_R"]
    return {k: round(float(v), 6) for k, v in zip(keys, raw)}


def calibrate_observer() -> dict[str, Any]:
    ensure_dirs()
    cfg = read_config()
    seed = int(cfg["seed"])
    rng = np.random.default_rng(seed)
    df = load_conflicts()
    val, test = split_eval(df)
    manual = {"weights": {"w_p": 0.35, "w_u": 0.15, "w_I": 0.20, "w_Delta": 0.10, "w_R": 0.20}, "thresholds": [0.25, 0.45, 0.65, 0.82], "gamma_max": 0.5, "I_min": 0.4, "Delta_max": 0.35}
    candidates: list[dict[str, Any]] = [manual]
    grid_weights = [
        [0.50, 0.10, 0.10, 0.10, 0.20],
        [0.35, 0.15, 0.20, 0.10, 0.20],
        [0.25, 0.15, 0.25, 0.10, 0.25],
        [0.20, 0.20, 0.25, 0.10, 0.25],
    ]
    grid_thetas = [[0.20, 0.40, 0.60, 0.80], [0.25, 0.45, 0.65, 0.82], [0.30, 0.50, 0.68, 0.85]]
    for gw in grid_weights:
        for th in grid_thetas:
            candidates.append({"weights": normalize_weights(np.array(gw)), "thresholds": th, "gamma_max": 0.5, "I_min": 0.4, "Delta_max": 0.35, "search": "grid_search"})
    for _ in range(160):
        raw = rng.dirichlet(np.ones(5))
        th = sorted(rng.uniform(0.12, 0.90, size=4))
        candidates.append({"weights": normalize_weights(raw), "thresholds": [round(float(x), 6) for x in th], "gamma_max": round(float(rng.uniform(0.25, 0.75)), 6), "I_min": round(float(rng.uniform(0.2, 0.7)), 6), "Delta_max": round(float(rng.uniform(0.15, 0.6)), 6), "search": "random_search"})
    scored = []
    for c in candidates:
        m = objective(val, apply_config(val, c))
        scored.append((m["objective_J"], c, m))
    scored.sort(key=lambda x: x[0])
    best = scored[0][1]
    best_val = scored[0][2]
    best_grid = min((x for x in scored if x[1].get("search") == "grid_search"), key=lambda x: x[0])
    best_random = min((x for x in scored if x[1].get("search") == "random_search"), key=lambda x: x[0])
    best_test = objective(test, apply_config(test, best))
    manual_val = objective(val, apply_config(val, manual))
    data = {
        "data_split": {"validation_cases": int(len(val)), "test_cases": int(len(test)), "n_unique_objects": int(df["source_row_id"].nunique())},
        "search_space": {"methods": ["grid_search", "random_search"], "n_candidates": len(candidates)},
        "objective": "proxy-objective: 5*missed_critical + 3*false_auto_accept + 2*unsafe_accept_with_conflict + false_block + 0.5*excessive_defer",
        "best_config": best,
        "best_grid_config": best_grid[1],
        "best_random_config": best_random[1],
        "manual_config_metrics": manual_val,
        "best_grid_config_metrics_validation": best_grid[2],
        "best_random_config_metrics_validation": best_random[2],
        "best_config_metrics_validation": best_val,
        "best_config_metrics_test": best_test,
        "bootstrap_ci": {},
        "seed": seed,
    }
    (REPORT / "observer_calibration_report.json").write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    (CONFIG_DIR / "best_observer_config.yaml").write_text(dump_yaml(best), encoding="utf-8")
    write_calibration_md(data, manual, manual_val)
    return data


def write_calibration_md(data: dict[str, Any], manual: dict[str, Any], manual_metrics: dict[str, Any]) -> None:
    best = data["best_config"]
    bestm = data["best_config_metrics_validation"]
    best_grid = data["best_grid_config"]
    best_grid_m = data["best_grid_config_metrics_validation"]
    best_random = data["best_random_config"]
    best_random_m = data["best_random_config_metrics_validation"]
    def row(name: str, cfg: dict[str, Any], met: dict[str, Any]) -> str:
        w, th = cfg["weights"], cfg["thresholds"]
        return f"| {name} | {w['w_p']:.3f} | {w['w_u']:.3f} | {w['w_I']:.3f} | {w['w_Delta']:.3f} | {w['w_R']:.3f} | {th[0]:.3f} | {th[1]:.3f} | {th[2]:.3f} | {th[3]:.3f} | {met['missed_critical_ruptures']} | {met['false_auto_accept']} | {met['false_block']} | {met['objective_J']:.2f} |"
    md = """# Автоматическая калибровка риск-ориентированного наблюдателя

Калибровка выполнена по proxy-objective, а не по клинической экспертной разметке.

`rho = w_p*rho_pred + w_u*u_M + w_I*(1-I_pre) + w_Delta*Delta_M + w_R*chi_R`.

Таблица 3.X. Автоматическая калибровка риск-ориентированного наблюдателя.

| Режим | w_p | w_u | w_I | w_Delta | w_R | theta_1 | theta_2 | theta_3 | theta_4 | missed critical | false auto accept | false block | objective J |
| --- | ---: | ---: | ---: | ---: | ---: | ---: | ---: | ---: | ---: | ---: | ---: | ---: | ---: |
"""
    md += row("manual_config", manual, manual_metrics) + "\n"
    md += row("best_grid_config", best_grid, best_grid_m) + "\n"
    md += row("best_random_config", best_random, best_random_m) + "\n"
    md += row("best_config", best, bestm) + "\n"
    md += "\nВыбранные параметры сохраняются в ExplainPlan и trace.\n"
    (REPORT / "observer_calibration_report.md").write_text(md, encoding="utf-8")


def bootstrap_ci(values: np.ndarray, func, b: int = 2000, seed: int = 42) -> list[float]:
    rng = np.random.default_rng(seed)
    n = len(values)
    stats = []
    for _ in range(b):
        idx = rng.integers(0, n, size=n)
        stats.append(float(func(values[idx])))
    return [float(np.percentile(stats, 2.5)), float(np.percentile(stats, 97.5))]


def make_tables_and_patches() -> None:
    ensure_dirs()
    cfg = read_config()
    df = load_conflicts()
    diff = (df["action_f0"] != df["action_nas"]).astype(int).to_numpy()
    accept_block = ((df["action_f0"] == "accept") & (df["action_nas"] == "block")).astype(int).to_numpy()
    critical = df["chi_R_crit"].astype(int).to_numpy()
    detected = ((df["chi_R_crit"] == 1) & (df["action_nas"] == "block")).astype(int).to_numpy()
    false_auto = ((df["action_nas"] == "accept") & ((df["y_true"] != df["y_pred"]) | (df["chi_R"] == 1))).astype(int).to_numpy()
    calib = json.loads((REPORT / "observer_calibration_report.json").read_text(encoding="utf-8")) if (REPORT / "observer_calibration_report.json").exists() else calibrate_observer()
    val, _ = split_eval(df)
    best_actions = apply_config(val, calib["best_config"])
    j_vals = []
    for _, row in val.iterrows():
        one = pd.DataFrame([row])
        j_vals.append(objective(one, apply_config(one, calib["best_config"]))["objective_J"])
    recall_vals = detected[critical == 1]
    ci = {
        "seed": int(cfg["seed"]),
        "B": int(cfg["bootstrap_samples"]),
        "n_cases": int(len(df)),
        "n_unique_objects": int(df["source_row_id"].nunique()),
        "recall_critical_rupture": {"estimate": float(recall_vals.mean()) if len(recall_vals) else 0.0, "ci95": bootstrap_ci(recall_vals, np.mean, int(cfg["bootstrap_samples"]), int(cfg["seed"])) if len(recall_vals) else [0, 0]},
        "false_auto_accept_rate": {"estimate": float(false_auto.mean()), "ci95": bootstrap_ci(false_auto, np.mean, int(cfg["bootstrap_samples"]), int(cfg["seed"]))},
        "action_difference_f0_vs_nas": {"estimate": float(diff.mean()), "ci95": bootstrap_ci(diff, np.mean, int(cfg["bootstrap_samples"]), int(cfg["seed"]))},
        "f0_accept_to_nas_block_rate": {"estimate": float(accept_block.mean()), "ci95": bootstrap_ci(accept_block, np.mean, int(cfg["bootstrap_samples"]), int(cfg["seed"]))},
        "calibration_objective_J": {"estimate": float(np.sum(j_vals)), "ci95": bootstrap_ci(np.array(j_vals), np.sum, int(cfg["bootstrap_samples"]), int(cfg["seed"]))},
    }
    (REPORT / "bootstrap_ci_report.json").write_text(json.dumps(ci, ensure_ascii=False, indent=2), encoding="utf-8")
    write_patches(df, ci, calib)


def pct(x: float) -> str:
    return f"{100*x:.2f}%"


def write_patches(df: pd.DataFrame, ci: dict[str, Any], calib: dict[str, Any]) -> None:
    summary = pd.read_csv(REPORT / "f0_vs_nas_action_diff.csv").iloc[0].to_dict() if (REPORT / "f0_vs_nas_action_diff.csv").exists() else f0_vs_nas().iloc[0].to_dict()
    PATCH.mkdir(exist_ok=True)
    (PATCH / "chapter3_remove_hesitant_from_core.md").write_text("""# Замена по разделам 3.3, 3.6, 3.7, 3.8, 3.9, 3.14, 3.15, 3.22

Минимальное ядро главы 3 задаётся как `F_core = {F0, F_int, NAS, F_ML}`, где `NAS = F_N^src`. Класс `F_H` исключается из основного ядра и не используется в финальной экспериментальной проверке. Hesitant-представления могут рассматриваться как внешний расширяемый класс, но в минимальное ядро главы 3 не входят, поскольку не участвуют в финальной экспериментальной проверке.

Экспертное расхождение `u_expert`, если оно возникает, интерпретируется как частный случай источникового конфликта в `NAS` либо выносится в электронное приложение.
""", encoding="utf-8")
    (PATCH / "chapter3_real_conflict_experiment_insert.md").write_text(f"""# Вставка: эксперимент на реальных объяснительных конфликтах

Эксперимент выполнен на Breast Cancer Wisconsin. Это не клиническая апробация, а проверка XAI-конфликтов на реальной табличной задаче. В датасете 569 уникальных объектов; 1000 evaluation cases сформированы bootstrap-выборкой из validation/test split. Поэтому результат описывает evaluation cases, а не уникальные клинические объекты.

Real explanation conflict определяется как объединение rank conflict, sign conflict и rule conflict между SHAP, LIME и rule/ExplainPlan-профилем. Используются `k=5` и `theta_rank=0.4`.

| Показатель | Значение |
| --- | ---: |
| cases | {int(summary['n_cases'])} |
| unique objects | {int(summary['n_unique_objects'])} |
| conflicts | {int(summary['n_conflicts'])} |
| critical conflicts | {int(summary['n_critical_conflicts'])} |
| action diff F0 vs NAS | {pct(float(summary['share_action_diff_f0_nas']))} [{pct(float(summary['ci95_action_diff_low']))}; {pct(float(summary['ci95_action_diff_high']))}] |
| F0 accept -> NAS block | {pct(float(summary['share_f0_accept_nas_block']))} [{pct(float(summary['ci95_f0_accept_nas_block_low']))}; {pct(float(summary['ci95_f0_accept_nas_block_high']))}] |

В режиме `F0` конфликт источников редуцируется к одной степени принадлежности. В режиме `NAS/F_ML` конфликт сохраняется как отдельная компонента источнико-аннотированного представления.
""", encoding="utf-8")
    best = calib["best_config"]
    (PATCH / "chapter3_auto_calibration_insert.md").write_text(f"""# Вставка: автоматическая калибровка наблюдателя

Калибровка выполнена по proxy-objective, а не по клинической экспертной разметке.

`rho = w_p*rho_pred + w_u*u_M + w_I*(1-I_pre) + w_Delta*Delta_M + w_R*chi_R`.

Целевая функция:

`J = 5*missed_critical_ruptures + 3*false_auto_accept + 2*unsafe_accept_with_conflict + false_block + 0.5*excessive_defer`.

Лучшие параметры сохранены в `configs/chapter3/best_observer_config.yaml`:

```yaml
{dump_yaml(best)}
```

Выбранные параметры должны фиксироваться в `ExplainPlan` и trace маршрута.
""", encoding="utf-8")
    (PATCH / "chapter3_updated_defended_positions.md").write_text("""# Обновлённые защищаемые положения главы 3

1. **Невыразимость и минимальное ядро.** Показано, что класс F0 не различает часть XAI-разрывов, а минимальное рабочее ядро `{F0, F_int, NAS, F_ML}` покрывает типы неопределённости, реально используемые в наблюдающем контуре.

2. **Редукция и действие.** Для перехода от более выразительного представления к пользовательской форме вычисляется потеря `Delta`, которая входит в `d_E^ext` и риск-наблюдатель.

3. **Калибруемый риск-наблюдатель.** Действие наблюдателя определяется не вручную, а через формулу риска, таблицу порогов и калибруемые параметры ExplainPlan; критический разрыв `chi_R^crit` блокирует автоматическое применение независимо от числового риска.
""", encoding="utf-8")


def validate_package() -> None:
    ensure_dirs()
    required = [
        REPORT / "current_chapter_audit.md",
        REPORT / "hesitant_dependency_map.csv",
        REPORT / "hesitant_removal_plan.md",
        REPORT / "real_conflict_experiment_report.md",
        REPORT / "real_conflict_summary.csv",
        REPORT / "f0_vs_nas_action_diff.csv",
        REPORT / "bootstrap_ci_report.json",
        REPORT / "observer_calibration_report.md",
        REPORT / "observer_calibration_report.json",
        CONFIG_DIR / "best_observer_config.yaml",
        PATCH / "chapter3_remove_hesitant_from_core.md",
        PATCH / "chapter3_real_conflict_experiment_insert.md",
        PATCH / "chapter3_auto_calibration_insert.md",
        PATCH / "chapter3_updated_defended_positions.md",
        ROOT / "README_chapter3_reproduce.md",
    ]
    lines = ["# Валидация пакета доказательных артефактов главы 3\n"]
    ok = True
    for p in required:
        exists = p.exists()
        ok = ok and exists
        lines.append(f"- {'OK' if exists else 'FAIL'} `{p.relative_to(ROOT)}`")
    if (REPORT / "real_conflict_summary.csv").exists():
        df = pd.read_csv(REPORT / "real_conflict_summary.csv")
        lines.append(f"- n_cases: {len(df)}")
        lines.append(f"- n_unique_objects: {df['source_row_id'].nunique()}")
        lines.append(f"- split present: {'OK' if 'split' in df and df['split'].notna().all() else 'FAIL'}")
        lines.append(f"- seed present: {'OK' if 'seed' in df else 'FAIL'}")
        text_all = "\n".join(p.read_text(encoding="utf-8", errors="ignore") for p in [REPORT / "real_conflict_experiment_report.md", PATCH / "chapter3_real_conflict_experiment_insert.md"] if p.exists())
        lines.append(f"- есть ограничение про не клиническую апробацию: {'OK' if 'не клиническая апробация' in text_all else 'FAIL'}")
        lines.append(f"- нет фразы `1000 пациентов`: {'OK' if '1000 пациентов' not in text_all else 'FAIL'}")
        lines.append(f"- явно указаны реальные конфликты: {'OK' if 'реальн' in text_all.lower() and 'SHAP' in text_all and 'LIME' in text_all else 'FAIL'}")
    if (REPORT / "f0_vs_nas_action_diff.csv").exists():
        s = pd.read_csv(REPORT / "f0_vs_nas_action_diff.csv").iloc[0]
        lines.append(f"- фактический процент action diff из CSV: {100*float(s['share_action_diff_f0_nas']):.2f}%")
    lines.append("\nИтог: " + ("PASS" if ok else "FAIL"))
    (REPORT / "package_validation_report.md").write_text("\n".join(lines) + "\n", encoding="utf-8")
    make_zip()


def make_zip() -> None:
    files = [
        "reports/chapter3/current_chapter_audit.md",
        "reports/chapter3/hesitant_dependency_map.csv",
        "reports/chapter3/hesitant_removal_plan.md",
        "reports/chapter3/real_conflict_experiment_report.md",
        "reports/chapter3/real_conflict_summary.csv",
        "reports/chapter3/f0_vs_nas_action_diff.csv",
        "reports/chapter3/bootstrap_ci_report.json",
        "reports/chapter3/observer_calibration_report.md",
        "reports/chapter3/observer_calibration_report.json",
        "configs/chapter3/best_observer_config.yaml",
        "patches/chapter3_remove_hesitant_from_core.md",
        "patches/chapter3_real_conflict_experiment_insert.md",
        "patches/chapter3_auto_calibration_insert.md",
        "patches/chapter3_updated_defended_positions.md",
        "scripts/chapter3_audit_docx.py",
        "scripts/chapter3_build_real_conflicts.py",
        "scripts/chapter3_f0_vs_nas_experiment.py",
        "scripts/chapter3_calibrate_observer.py",
        "scripts/chapter3_make_tables.py",
        "scripts/chapter3_validate_package.py",
        "scripts/chapter3_evidence_common.py",
        "Makefile",
        "README_chapter3_reproduce.md",
    ]
    with zipfile.ZipFile(ROOT / "chapter3_final_fix_evidence_package.zip", "w", zipfile.ZIP_DEFLATED) as z:
        for f in files:
            p = ROOT / f
            if p.exists():
                z.write(p, f)
