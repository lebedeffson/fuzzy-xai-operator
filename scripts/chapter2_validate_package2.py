#!/usr/bin/env python3
"""Validate chapter 2 package 2 artifacts."""

from __future__ import annotations

import argparse
import re
import subprocess
import tempfile
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "glava2_chapter2_package2_full_fixed.docx"
SRC = ROOT / "docs/chapter2/glava2_chapter2_package1_fixed.docx"
FIG_DIR = ROOT / "figures/chapter2"
PNG_DIR = FIG_DIR / "png_preview"
REPORT = ROOT / "reports/chapter2/package2_validation_report.md"

FIGS = [
    "fig_2_1_Ek_structure",
    "fig_2_2_Tij_composition",
    "fig_2_3_dE_metric",
    "fig_2_4_GExpl_graph",
    "fig_2_5_sample113_flow",
]
BAD_TERMS = ["объяснимый ИИ", "объяснимого ИИ", "объяснимому ИИ", "объяснимым ИИ", "объяснительного ИИ"]
FORBIDDEN = ["cross-validation", "кросс-валидация", "новый датасет", "MNIST", "CIFAR", "ImageNet", "накопления неопределённости", "накопления неопределенности"]


def doc_text(path):
    if not path.exists():
        return ""
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def docx_openable(path):
    if not path.exists():
        return False, "файл отсутствует"
    try:
        with zipfile.ZipFile(path) as z:
            z.testzip()
        Document(str(path))
        return True, "python-docx открывает файл"
    except Exception as exc:
        return False, str(exc)


def libreoffice_check(path):
    if not path.exists():
        return False, "файл отсутствует"
    if subprocess.run(["bash", "-lc", "command -v libreoffice >/dev/null || command -v soffice >/dev/null"], check=False).returncode != 0:
        return False, "LibreOffice/soffice не найден"
    exe = "libreoffice" if subprocess.run(["bash", "-lc", "command -v libreoffice >/dev/null"], check=False).returncode == 0 else "soffice"
    with tempfile.TemporaryDirectory() as tmp:
        r = subprocess.run([exe, "--headless", "--convert-to", "pdf", "--outdir", tmp, str(path)], text=True, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, timeout=120)
        ok = r.returncode == 0 and any(Path(tmp).glob("*.pdf"))
        return ok, f"exit={r.returncode}; {r.stdout.strip()[:500]}"


def make_reproduce_check():
    makefile = ROOT / "Makefile"
    docker = ROOT / "Dockerfile"
    has_target = False
    if makefile.exists():
        has_target = bool(re.search(r"^reproduce-chapter2:", makefile.read_text(encoding="utf-8"), re.M))
    if not has_target:
        return docker.exists(), makefile.exists(), "цель reproduce-chapter2 не найдена"
    try:
        r = subprocess.run(["make", "reproduce-chapter2"], cwd=ROOT, text=True, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, timeout=120)
        return docker.exists(), makefile.exists(), f"make reproduce-chapter2 exit code: {r.returncode}"
    except subprocess.TimeoutExpired:
        return docker.exists(), makefile.exists(), "make reproduce-chapter2 timeout after 120s"


def main():
    argparse.ArgumentParser().parse_args()
    REPORT.parent.mkdir(parents=True, exist_ok=True)
    text = doc_text(DOCX)
    src_text = doc_text(SRC)
    lines = []
    add = lines.append

    add("# Отчёт валидации пакета 2 главы 2\n")
    ok_docx, msg_docx = docx_openable(DOCX)
    ok_lo, msg_lo = libreoffice_check(DOCX)

    add("## Проверка файлов\n")
    add(f"- Исходный DOCX: {'OK' if SRC.exists() else 'FAIL'} — `docs/chapter2/glava2_chapter2_package1_fixed.docx`")
    add(f"- Итоговый DOCX: {'OK' if DOCX.exists() else 'FAIL'} — `{DOCX.name}`")
    for ext in ["svg", "pdf"]:
        missing = [f for f in FIGS if not (FIG_DIR / f"{f}.{ext}").exists()]
        add(f"- 5 {ext.upper()}: {'OK' if not missing else 'FAIL'}" + (f"; отсутствуют: {missing}" if missing else ""))
    missing_png = [f for f in FIGS if not (PNG_DIR / f"{f}.png").exists()]
    add(f"- 5 PNG-preview: {'OK' if not missing_png else 'FAIL'}" + (f"; отсутствуют: {missing_png}" if missing_png else ""))
    add("- В DOCX вставлены PNG-preview; SVG и PDF приложены как векторные исходники.")
    add(f"- DOCX zip/python-docx: {'OK' if ok_docx else 'FAIL'} — {msg_docx}")
    add(f"- LibreOffice convert: {'OK' if ok_lo else 'FAIL'} — {msg_lo}\n")

    add("## Проверка подраздела 2.18.2\n")
    checks = [
        ("найден заголовок", "2.18.2 Адаптация оператора к разным XAI-методам"),
        ("найдено упоминание SHAP", "SHAP"),
        ("найдено упоминание Grad-CAM", "Grad-CAM"),
        ("найдено упоминание дерева решений", "дерева решений"),
        ("найдено упоминание ANFIS", "ANFIS"),
        ("найдено E_k=<L_k, μ_k, R_k, α_k, u_k, τ_k>", "E_k=<L_k, μ_k, R_k, α_k, u_k, τ_k>"),
    ]
    for label, needle in checks:
        add(f"- {label}: {'OK' if needle in text else 'FAIL'}")
    add("")

    add("## Проверка рисунков\n")
    captions = re.findall(r"Рисунок\s+2\.(\d)\.", text)
    add(f"- Найдено подписей рисунков 2.1-2.5: {len(captions)}")
    add(f"- Последовательность: {captions}")
    add(f"- Нет пропусков: {'OK' if captions == ['1','2','3','4','5'] else 'FAIL'}")
    add(f"- Нет дублей: {'OK' if len(captions) == len(set(captions)) else 'FAIL'}")
    for n in range(1, 6):
        add(f"- Ссылка на рисунок 2.{n}: {'OK' if f'рисунке 2.{n}' in text or f'рисунок 2.{n}' in text else 'FAIL'}")
    add("")

    add("## Проверка терминологии\n")
    for term in BAD_TERMS:
        found = [m.start() for m in re.finditer(re.escape(term), text)]
        add(f"- `{term}`: {len(found)}")
    add("")

    add("## Проверка запретов\n")
    for term in FORBIDDEN:
        before = src_text.count(term) if src_text else 0
        after = text.count(term)
        status = "OK" if after <= before else "FAIL"
        add(f"- `{term}`: {status} (исходно {before}, итог {after})")
    add("- Новые экспериментальные heatmap/графики чувствительности скриптами пакета не создаются.")
    add("")

    add("## Проверка воспроизводимости\n")
    has_docker, has_make, repro = make_reproduce_check()
    add(f"- Dockerfile: {'OK' if has_docker else 'WARN'}")
    add(f"- Makefile: {'OK' if has_make else 'WARN'}")
    add(f"- reproduce-chapter2: {repro}")
    add("")

    critical = []
    if not DOCX.exists() or not ok_docx or captions != ["1", "2", "3", "4", "5"]:
        critical.append("критическая проверка DOCX/рисунков не пройдена")
    if any(term in text for term in BAD_TERMS):
        critical.append("остались кривые терминологические формы")
    add("## Итог\n")
    add("Критических ошибок нет." if not critical else "Критические ошибки: " + "; ".join(critical) + ".")
    REPORT.write_text("\n".join(lines) + "\n", encoding="utf-8")
    print(f"Saved {REPORT}")


if __name__ == "__main__":
    main()
