#!/usr/bin/env python3
from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/chapters/glava_4_FuzzyXAI_corrected_final.docx"
OUTPUT = ROOT / "docs/chapters/glava_4_FuzzyXAI_h10_c4_revision.docx"


def _rows(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def _paragraph_by_prefix(document: Document, prefix: str):
    return next(item for item in document.paragraphs if item.text.startswith(prefix))


def _set_text(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def _insert_paragraph(anchor, text: str, *, bold: bool = False, code: bool = False):
    paragraph = OxmlElement("w:p")
    anchor._p.addprevious(paragraph)
    from docx.text.paragraph import Paragraph

    wrapped = Paragraph(paragraph, anchor._parent)
    run = wrapped.add_run(text)
    run.bold = bold
    if code:
        run.font.name = "Courier New"
        run.font.size = Pt(9)
    return wrapped


def _insert_page_break(anchor) -> None:
    paragraph = _insert_paragraph(anchor, "")
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def _insert_table(
    document: Document,
    anchor,
    headers: tuple[str, ...],
    rows: list[tuple[str, ...]],
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
    anchor._p.addprevious(table._tbl)


def build() -> Path:
    document = Document(SOURCE)
    _set_text(
        _paragraph_by_prefix(document, "4.8."),
        "4.8. Технические представления диагностического состояния",
    )
    _set_text(
        document.paragraphs[53],
        "Реализованы три технических представления одного диагностического "
        "состояния: пользовательское, инженерное и аудиторское. Интерфейс не "
        "вычисляет диагностические состояния, стоимость разреза или результат "
        "повторной сертификации, а отображает значения проверенного серверного "
        "пакета.",
    )
    _set_text(
        document.paragraphs[56],
        "Понятность представлений, влияние интерфейса на скорость локализации "
        "и соответствие рабочим практикам специалистов в настоящей работе не "
        "оценивались. Технические сведения о панелях и автоматических проверках "
        "вынесены в supplementary-материалы.",
    )
    _set_text(
        _paragraph_by_prefix(document, "4.9."),
        "4.9. H10-C3 R4: вычислительная проверка на заранее сгенерированных "
        "контролируемых структурных мутациях",
    )
    h10_intro = document.paragraphs[58]
    _set_text(
        h10_intro,
        h10_intro.text.replace(
            "из 240 независимых структурных шаблонов",
            "из 240 независимых заранее сгенерированных контролируемых "
            "структурных мутаций",
        ),
    )
    _set_text(
        document.paragraphs[62],
        "Полученный вывод подтверждён только в пределах заранее "
        "сгенерированных контролируемых структурных мутаций и относится к "
        "алгоритмической диагностике объяснительного маршрута, поиску "
        "оптимального диагностического разреза и полной повторной "
        "сертификации. Реальные программные инциденты, понятность, экспертная "
        "полезность, промышленная или клиническая безопасность и "
        "организационная эффективность не оценивались.",
    )

    anchor = _paragraph_by_prefix(document, "4.10.")
    bootstrap = _rows(ROOT / "results/h10_c4/BOOTSTRAP_INTERVALS.csv")
    descriptive = _rows(ROOT / "results/h10_c4/STRATEGY_COMPARISON.csv")
    status = json.loads(
        (ROOT / "results/h10_c4/H10_C4_FINAL_STATUS.json").read_text()
    )

    _insert_page_break(anchor)
    _insert_paragraph(
        anchor,
        "4.10. Операционная стоимость восстановления на контролируемых мутациях",
        bold=True,
    )
    _insert_paragraph(
        anchor,
        "H10-C4 оформлен как отдельный перспективный протокол и не изменяет "
        "закрытый результат H10-C3. На 120 новых отложенных контролируемых "
        "структурных мутациях шести семейств сравнивались repair-all (B_ALL), "
        "первый валидный план (B_FIRST), локально-жадная стратегия (B_GREEDY) "
        "и глобальный минимальный разрез (O_GLOBAL). Все стратегии начинали с "
        "одного снимка состояния, использовали одинаковые контракты, "
        "разрешённые операции, проверяющий модуль и полную повторную "
        "сертификацию.",
    )
    _insert_table(
        document,
        anchor,
        ("Стратегия", "Правило выбора"),
        [
            ("B_ALL", "Исправить все объекты зарегистрированных нарушений"),
            ("B_FIRST", "Первый найденный валидный план"),
            ("B_GREEDY", "Локально минимальная стоимость на шаге"),
            ("O_GLOBAL", "Глобальный минимальный диагностический разрез"),
        ],
    )
    _insert_paragraph(anchor, "Таблица 4.8. Стратегии восстановления H10-C4.")
    _insert_paragraph(
        anchor,
        "Глобальный разрез обеспечил полную повторную сертификацию в 120 из "
        "120 сценариев и не создал новых критических нарушений. Средняя "
        "нормализованная исполняемая стоимость O_GLOBAL составила 0,516 при "
        "1,000 для B_ALL и B_FIRST и 0,594 для B_GREEDY. Среднее число "
        "действий уменьшилось с 3,10 для B_ALL до 1,00, число затронутых "
        "компонентов — с 3,10 до 1,00, число повторных проверок — с 14,30 до "
        "8,00. Это алгоритмическая экономия операций; время инженера не "
        "измерялось.",
    )
    comparison_rows = []
    for row in bootstrap:
        ci = json.loads(row["ci_95"])
        comparison_rows.append(
            (
                row["comparison"].replace("_vs_", " − "),
                f"{float(row['mean_difference']):.6f}",
                f"[{ci[0]:.6f}; {ci[1]:.6f}]",
                f"{float(row['holm_p']):.6f}",
            )
        )
    _insert_table(
        document,
        anchor,
        ("Сравнение", "Δ стоимости", "95%-й ДИ", "p Холма"),
        comparison_rows,
    )
    _insert_paragraph(
        anchor,
        "Таблица 4.9. Парные сравнения нормализованной исполняемой стоимости; "
        "отрицательная разность O_GLOBAL − baseline соответствует преимуществу "
        "глобального разреза.",
    )
    _insert_paragraph(
        anchor,
        "Для всех трёх сравнений общий парный bootstrap-поток включал 10 000 "
        "итераций; SHA256 индексного потока равен "
        f"{bootstrap[0]['bootstrap_index_stream_sha256']}. Все доверительные "
        "интервалы исключали ноль после заранее зарегистрированной поправки "
        "Холма.",
    )

    def mean(strategy: str, metric: str) -> float:
        row = next(
            item
            for item in descriptive
            if item["strategy"] == strategy and item["metric"] == metric
        )
        return float(row["mean"])

    _insert_table(
        document,
        anchor,
        ("Стратегия", "Действия", "Компоненты", "Проверки", "Машинное время, мс"),
        [
            (
                strategy,
                f"{mean(strategy, 'repair_action_count'):.2f}",
                f"{mean(strategy, 'unique_touched_components'):.2f}",
                f"{mean(strategy, 'recertification_check_count'):.2f}",
                f"{mean(strategy, 'execution_time_ms'):.3f}",
            )
            for strategy in ("B_ALL", "B_FIRST", "B_GREEDY", "O_GLOBAL")
        ],
    )
    _insert_paragraph(
        anchor,
        "Таблица 4.10. Фактически исполненные операции. Машинное время не "
        "интерпретируется как время работы инженера.",
    )
    _insert_table(
        document,
        anchor,
        ("Модель", "Величина", "Интерпретация", "Ограничение"),
        [
            ("Uniform", "Действия", "Структурная", "Не различает операции"),
            ("Runtime", "Время CPU", "Вычислительная", "Не время инженера"),
            ("Dependency", "Fan-out", "Структурный охват", "Не орг. затраты"),
            ("Hybrid", "Сумма факторов", "Чувствительность", "Зависит от весов"),
        ],
    )
    _insert_paragraph(anchor, "Таблица 4.11. Зарегистрированные модели стоимости.")
    _insert_paragraph(
        anchor,
        "В сетке из 48 заранее зарегистрированных комбинаций весов 108 из "
        "120 сценариев достигли порога устойчивости 80%; доля таких сценариев "
        f"составила {status['selection_stability_rate']:.3f}. Остальные 12 "
        "сценариев сохранены в отчёте как чувствительные к способу задания "
        "стоимости, а не исключены из анализа.",
    )
    _insert_paragraph(anchor, "4.10.1. Контрактная композиция операторов", bold=True)
    _insert_paragraph(
        anchor,
        "Исполняемый интерфейс задаёт operator_id, входной и выходной "
        "контракты, apply, verify_preconditions и verify_postconditions. "
        "Несовместимость первого выходного и второго входного контрактов "
        "возвращает структурированную ошибку OPERATOR_CONTRACT_MISMATCH.",
    )
    _insert_paragraph(
        anchor,
        "class DiagnosticOperator(Protocol):\n"
        "    operator_id: str\n"
        "    input_contract: OperatorContract\n"
        "    output_contract: OperatorContract\n"
        "    def apply(self, state: RouteState) -> RouteState: ...\n"
        "    def verify_preconditions(self, state: RouteState) -> VerificationResult: ...\n"
        "    def verify_postconditions(self, state: RouteState) -> VerificationResult: ...",
        code=True,
    )
    _insert_paragraph(
        anchor,
        "O_diag = O_recert ◦ O_repair ◦ O_cut ◦ O_contract ◦ O_provenance",
        code=True,
    )
    _insert_paragraph(
        anchor,
        "Реализована последовательная контрактная композиция исполняемых "
        "операторов. SHAP- и LIME-артефакты проходят общий оператор проверки "
        "происхождения, но их значения не объединяются математически. Полная "
        "символическая операторная алгебра с доказательством замкнутости, "
        "тождественными и обратными операторами для произвольных "
        "объяснительных объектов в текущей версии не реализована.",
    )

    _set_text(anchor, "4.11. Границы результата главы")
    conclusions = _paragraph_by_prefix(document, "4.11. Выводы")
    _set_text(conclusions, "4.12. Выводы по главе")
    h10_conclusion = next(
        item
        for item in document.paragraphs
        if item.text.startswith("На независимой закрытой совокупности H10-C3")
    )
    _set_text(
        h10_conclusion,
        "На независимой закрытой совокупности H10-C3 R4 подтверждены только в "
        "пределах заранее сгенерированных контролируемых структурных мутаций "
        "преимущества типизированного графового метода по принадлежности "
        "диагностического разреза множеству оптимальных решений и полной "
        "повторной сертификации; автоматический статус — SCIENTIFIC_PASS.",
    )
    closing = document.paragraphs[-1]
    _insert_paragraph(
        closing,
        "Отдельный перспективный эксперимент H10-C4 показал операционное "
        "преимущество глобального минимального разреза на 120 новых "
        "контролируемых структурных мутациях: уменьшились исполняемая "
        "стоимость, число действий, затронутых компонентов и повторных "
        "проверок при сохранении полной сертификации. Результат не "
        "распространяется на реальные инциденты или трудозатраты инженера.",
    )
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
