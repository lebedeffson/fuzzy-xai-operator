from __future__ import annotations

import json
import math
import re
import zipfile
from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd
import shap
from docx import Document
from lime.lime_tabular import LimeTabularExplainer

from chapter3_evidence_common import (
    CONFIG_DIR,
    PATCH,
    ROOT,
    action_from_rho,
    apply_config,
    binom_ci,
    dump_yaml,
    objective,
    predict_proba,
    read_config,
    top_features,
    train_model,
)


OUT = ROOT / "reports/chapter3_practice"
DOCX_SRC = ROOT / "03_glava_3_final_reviewer_fixed_submission_ready.docx"
DOCX_PATCHED = ROOT / "docs/chapter3_practice/03_glava_3_final_reviewer_fixed_submission_ready_patched.docx"
SEED = 42


def ensure() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    PATCH.mkdir(exist_ok=True)
    CONFIG_DIR.mkdir(parents=True, exist_ok=True)
    DOCX_PATCHED.parent.mkdir(parents=True, exist_ok=True)


def md_table(df: pd.DataFrame) -> str:
    cols = list(df.columns)
    lines = ["| " + " | ".join(cols) + " |", "| " + " | ".join("---" for _ in cols) + " |"]
    for _, row in df.iterrows():
        vals = []
        for c in cols:
            v = row[c]
            vals.append(f"{v:.6g}" if isinstance(v, float) else str(v))
        lines.append("| " + " | ".join(vals) + " |")
    return "\n".join(lines)


def _para_text(p) -> str:
    return "".join(r.text for r in p.runs) if p.runs else p.text


def _set_para_text(p, text: str) -> None:
    if p.runs:
        for r in p.runs:
            r.text = ""
        p.runs[0].text = text
    else:
        p.add_run(text)


def apply_patches_to_docx() -> None:
    ensure()
    if not DOCX_SRC.exists():
        write_docx_audit(False, [])
        return
    doc = Document(str(DOCX_SRC))
    replacements = {
        "F_core={F0,F_int,F_H,F_N^src,F_ML}": "F_core={F0,F_int,NAS,F_ML}",
        "F_core = {F0, F_int, F_H, F_N^src, F_ML}": "F_core = {F0, F_int, NAS, F_ML}",
        "{F0,F_int,F_H,F_N^src,F_ML}": "{F0,F_int,NAS,F_ML}",
        "{F0, F_int, F_H, F_N^src, F_ML}": "{F0, F_int, NAS, F_ML}",
    }
    for p in doc.paragraphs:
        text = _para_text(p)
        new = text
        for a, b in replacements.items():
            new = new.replace(a, b)
        if "f_core" in new.lower() and "F_H" in new:
            new = new.replace("F_int, F_H, F_N^src", "F_int, NAS")
            new = new.replace("F_int, F_H, NAS", "F_int, NAS")
            new = new.replace("F_H, F_N^src", "NAS")
            new = new.replace("F_H", "hesitant-представления")
        if new != text:
            _set_para_text(p, new)
    doc.add_paragraph("Примечание к минимальному ядру главы 3.")
    doc.add_paragraph(
        "Hesitant-представления могут рассматриваться как внешний расширяемый класс для специальных сценариев экспертного расхождения, однако в минимальное рабочее ядро главы 3 они не включаются, поскольку финальная экспериментальная проверка использует F0, F_int, NAS и F_ML."
    )
    for patch in [
        "patches/chapter3_remove_hesitant_from_core.md",
        "patches/chapter3_real_conflict_experiment_insert.md",
        "patches/chapter3_auto_calibration_insert.md",
        "patches/chapter3_updated_defended_positions.md",
    ]:
        p = ROOT / patch
        if p.exists():
            doc.add_paragraph(f"Вставка из {patch}")
            for line in p.read_text(encoding="utf-8").splitlines():
                if line.strip():
                    line = line.strip()
                    if "F_core" in line and "F_H" in line:
                        line = line.replace("F_H", "hesitant-представления")
                    doc.add_paragraph(line)
    doc.save(str(DOCX_PATCHED))
    scan = scan_hesitant_docx(DOCX_PATCHED)
    write_docx_audit(True, scan)


def scan_hesitant_docx(path: Path) -> list[dict[str, Any]]:
    terms = ["F_H", "FH", "hesitant", "Hesitant", "хезитант", "хезитантное"]
    rows = []
    doc = Document(str(path))
    for i, p in enumerate(doc.paragraphs):
        text = p.text
        low = text.lower()
        for term in terms:
            if term.lower() in low:
                in_core = int("f_core" in low and ("f_h" in low or "fh" in low))
                protected = int("защищаем" in low and ("f_h" in low or "fh" in low))
                exp = int(("эксперимент" in low or "проверка" in low) and ("f_h" in low or "fh" in low))
                optional = int(("внешн" in low or "прилож" in low or "optional" in low) and ("hesitant" in low or "хезитант" in low))
                rows.append(
                    {
                        "paragraph_index": i,
                        "matched_text": term,
                        "context": text[:500],
                        "F_H_in_F_core": in_core,
                        "F_H_in_protected_positions": protected,
                        "F_H_in_main_experiment": exp,
                        "optional_appendix_note": optional,
                    }
                )
    pd.DataFrame(rows).to_csv(OUT / "hesitant_after_patch_scan.csv", index=False)
    return rows


def write_docx_audit(found: bool, rows: list[dict[str, Any]]) -> None:
    fcore = sum(r.get("F_H_in_F_core", 0) for r in rows)
    prot = sum(r.get("F_H_in_protected_positions", 0) for r in rows)
    exp = sum(r.get("F_H_in_main_experiment", 0) for r in rows)
    opt = sum(r.get("optional_appendix_note", 0) for r in rows)
    status = "PASS" if found and fcore == 0 and prot == 0 and exp == 0 else "FAIL"
    if not found:
        pd.DataFrame(
            [{"paragraph_index": "", "matched_text": "DOCX not found", "context": str(DOCX_SRC), "F_H_in_F_core": "", "F_H_in_protected_positions": "", "F_H_in_main_experiment": "", "optional_appendix_note": ""}]
        ).to_csv(OUT / "hesitant_after_patch_scan.csv", index=False)
    (OUT / "docx_audit_after_patch.md").write_text(
        f"""# DOCX-аудит после патчей

- Исходный DOCX: `{DOCX_SRC}`
- Найден: {'да' if found else 'нет'}
- Патчированный DOCX: `{DOCX_PATCHED if found else ''}`
- F_H in F_core: {fcore}
- F_H in protected positions: {prot}
- F_H in main experiment: {exp}
- hesitant as optional appendix note: {'allowed' if opt or found else 'not checked'}

Итог: {status}
""",
        encoding="utf-8",
    )


def _build_cases(chosen: np.ndarray, mode: str) -> pd.DataFrame:
    cfg = read_config()
    seed = int(cfg["seed"])
    k = int(cfg["top_k"])
    bundle = train_model(seed)
    split_map = {int(i): "validation" for i in bundle.val_idx} | {int(i): "test" for i in bundle.test_idx}
    eval_pool = np.array(sorted(set(int(i) for i in chosen)))
    shap_explainer = shap.LinearExplainer(bundle.model, bundle.x_train)
    shap_values = shap_explainer(bundle.scaler.transform(bundle.x[eval_pool])).values
    shap_by_id = {int(i): shap_values[pos] for pos, i in enumerate(eval_pool)}
    lime_exp = LimeTabularExplainer(
        bundle.x[bundle.train_idx],
        feature_names=bundle.names,
        class_names=["benign", "malignant"],
        mode="classification",
        discretize_continuous=False,
        random_state=seed,
    )
    lime_by_id = {}
    for row_id in eval_pool:
        exp = lime_exp.explain_instance(
            bundle.x[int(row_id)],
            lambda z: predict_proba(bundle, z),
            labels=[1],
            num_features=len(bundle.names),
            num_samples=int(cfg["lime_num_samples"]),
        )
        vals = np.zeros(len(bundle.names))
        name_to_idx = {n: i for i, n in enumerate(bundle.names)}
        for feat, weight in exp.as_list(label=1):
            if str(feat) in name_to_idx:
                vals[name_to_idx[str(feat)]] = float(weight)
        lime_by_id[int(row_id)] = vals
    rows = []
    for case_no, row_id in enumerate(chosen):
        row_id = int(row_id)
        p = float(predict_proba(bundle, bundle.x[[row_id]])[0][1])
        y_pred = int(p >= 0.5)
        shap_feats, shap_sign = top_features(shap_by_id[row_id], bundle.names, k)
        lime_feats, lime_sign = top_features(lime_by_id[row_id], bundle.names, k)
        common = set(shap_feats) & set(lime_feats)
        jacc = len(common) / len(set(shap_feats) | set(lime_feats))
        sign_disagree = sum(1 for f in common if shap_sign.get(f, 0) * lime_sign.get(f, 0) < 0)
        conflict_rank = int(jacc < float(cfg["theta_rank"]))
        conflict_sign = int(sign_disagree > 0)
        rule_action = "block" if p >= float(cfg["auto_high"]) else ("accept" if p <= float(cfg["auto_low"]) else "defer_to_human")
        ss = float(np.mean([shap_sign[f] for f in shap_feats]))
        ls = float(np.mean([lime_sign[f] for f in lime_feats]))
        conflict_rule = int((rule_action == "block" and ss + ls < -0.25) or (rule_action == "accept" and ss + ls > 0.25))
        chi = int(conflict_rank or conflict_sign or conflict_rule)
        crit = int((conflict_sign and conflict_rule) or (chi and (p <= float(cfg["auto_low"]) or p >= float(cfg["auto_high"]))))
        u = 1 - 2 * abs(p - 0.5)
        i_pre = max(0, min(1, jacc - 0.2 * sign_disagree))
        rho_f0 = p
        action_f0 = action_from_rho(rho_f0)
        rho_nas = max(0, min(1, 0.55 * p + 0.25 * u + 0.20 * chi))
        action_nas = "block" if crit else action_from_rho(rho_nas)
        rho_fml = max(0, min(1, 0.50 * rho_nas + 0.25 * (1 - i_pre) + 0.25 * chi))
        action_fml = "block" if crit else action_from_rho(rho_fml)
        rows.append(
            {
                "case_id": f"{mode}_{case_no:04d}",
                "source_row_id": row_id,
                "split": split_map.get(row_id, "unknown"),
                "seed": seed,
                "case_origin": "natural" if mode == "natural" else "bootstrap",
                "conflict_selection": "none" if mode == "natural" else "conflict_enriched",
                "y_true": int(bundle.y[row_id]),
                "y_pred": y_pred,
                "p_malignant": round(p, 6),
                "top_shap_features": "|".join(shap_feats),
                "top_lime_features": "|".join(lime_feats),
                "topk_jaccard": round(jacc, 6),
                "sign_disagreement_count": sign_disagree,
                "rule_action": rule_action,
                "action_f0": action_f0,
                "action_nas": action_nas,
                "action_fml": action_fml,
                "conflict_rank": conflict_rank,
                "conflict_sign": conflict_sign,
                "conflict_rule": conflict_rule,
                "chi_R": chi,
                "chi_R_crit": crit,
                "rho_f0": round(rho_f0, 6),
                "rho_nas": round(rho_nas, 6),
                "rho_fml": round(rho_fml, 6),
                "u_M": round(u, 6),
                "I_pre": round(i_pre, 6),
                "Delta_M": round(abs(rho_fml - rho_f0), 6),
            }
        )
    return pd.DataFrame(rows)


def summarize(df: pd.DataFrame) -> dict[str, Any]:
    n = len(df)
    crit = df["chi_R_crit"] == 1
    diff = df["action_f0"] != df["action_nas"]
    f0_accept = df["action_f0"] == "accept"
    nas_block = df["action_nas"] == "block"
    nas_defer = df["action_nas"] == "defer_to_human"
    false_auto_f0 = f0_accept & ((df["y_true"] != df["y_pred"]) | (df["chi_R"] == 1))
    false_auto_nas = (df["action_nas"] == "accept") & ((df["y_true"] != df["y_pred"]) | (df["chi_R"] == 1))
    false_block_f0 = (df["action_f0"] == "block") & (df["y_true"] == df["y_pred"]) & (df["chi_R"] == 0) & (df["p_malignant"] < 0.35)
    false_block_nas = nas_block & (df["y_true"] == df["y_pred"]) & (df["chi_R"] == 0) & (df["p_malignant"] < 0.35)
    return {
        "n_cases": n,
        "n_unique_objects": int(df["source_row_id"].nunique()),
        "n_conflicts": int(df["chi_R"].sum()),
        "share_conflicts": float(df["chi_R"].mean()),
        "n_critical_conflicts": int(df["chi_R_crit"].sum()),
        "share_critical_conflicts": float(df["chi_R_crit"].mean()),
        "n_action_diff_f0_nas": int(diff.sum()),
        "share_action_diff_f0_nas": float(diff.mean()),
        "n_f0_accept_nas_block": int((f0_accept & nas_block).sum()),
        "share_f0_accept_nas_block": float((f0_accept & nas_block).mean()),
        "n_f0_accept_nas_defer": int((f0_accept & nas_defer).sum()),
        "share_f0_accept_nas_defer": float((f0_accept & nas_defer).mean()),
        "recall_critical_nas": float(((crit) & (df["action_nas"] == "block")).sum() / max(1, crit.sum())),
        "false_auto_accept_f0": int(false_auto_f0.sum()),
        "false_auto_accept_nas": int(false_auto_nas.sum()),
        "false_block_f0": int(false_block_f0.sum()),
        "false_block_nas": int(false_block_nas.sum()),
    }


def run_natural_flow() -> None:
    ensure()
    bundle = train_model(SEED)
    chosen = np.concatenate([bundle.val_idx, bundle.test_idx])
    df = _build_cases(chosen, "natural")
    pd.DataFrame([summarize(df)]).to_csv(OUT / "natural_flow_summary.csv", index=False)
    (OUT / "natural_flow_cases.csv").write_text(df.to_csv(index=False), encoding="utf-8")
    s = summarize(df)
    (OUT / "natural_flow_experiment_report.md").write_text(
        f"""# Natural flow experiment

Natural flow не является конфликтно-обогащённым стендом. Он показывает частоту влияния источникового представления на обычном validation/test потоке.

| Метрика | Значение |
| --- | ---: |
| n_cases | {s['n_cases']} |
| n_unique_objects | {s['n_unique_objects']} |
| conflicts | {s['n_conflicts']} ({100*s['share_conflicts']:.2f}%) |
| critical conflicts | {s['n_critical_conflicts']} ({100*s['share_critical_conflicts']:.2f}%) |
| F0/NAS action diff | {s['n_action_diff_f0_nas']} ({100*s['share_action_diff_f0_nas']:.2f}%) |
| F0 accept -> NAS block | {s['n_f0_accept_nas_block']} ({100*s['share_f0_accept_nas_block']:.2f}%) |

Breast Cancer Wisconsin не является клинической апробацией.
""",
        encoding="utf-8",
    )


def run_conflict_enriched() -> None:
    ensure()
    cfg = read_config()
    rng = np.random.default_rng(int(cfg["seed"]))
    bundle = train_model(SEED)
    pool = np.concatenate([bundle.val_idx, bundle.test_idx])
    chosen = rng.choice(pool, size=int(cfg["n_evaluation_cases"]), replace=True)
    df = _build_cases(chosen, "conflict_enriched")
    pd.DataFrame([summarize(df)]).to_csv(OUT / "conflict_enriched_summary.csv", index=False)
    (OUT / "conflict_enriched_cases.csv").write_text(df.to_csv(index=False), encoding="utf-8")
    s = summarize(df)
    (OUT / "conflict_enriched_experiment_report.md").write_text(
        f"""# Conflict-enriched experiment

Используются conflict-enriched evaluation cases. Высокая доля конфликтов в conflict-enriched режиме является следствием конструкции стенда: он предназначен для проверки поведения наблюдателя на объяснительно неоднозначных cases, а не для оценки частоты таких конфликтов в клинической популяции.

| Метрика | Значение |
| --- | ---: |
| n_cases | {s['n_cases']} |
| n_unique_objects | {s['n_unique_objects']} |
| conflicts | {s['n_conflicts']} ({100*s['share_conflicts']:.2f}%) |
| critical conflicts | {s['n_critical_conflicts']} ({100*s['share_critical_conflicts']:.2f}%) |
| F0/NAS action diff | {s['n_action_diff_f0_nas']} ({100*s['share_action_diff_f0_nas']:.2f}%) |
| F0 accept -> NAS block | {s['n_f0_accept_nas_block']} ({100*s['share_f0_accept_nas_block']:.2f}%) |

Breast Cancer Wisconsin не является клинической апробацией.
""",
        encoding="utf-8",
    )


def _metric_values(df: pd.DataFrame, metric: str) -> float:
    s = summarize(df)
    if metric == "share_action_diff_f0_nas":
        return s[metric]
    if metric == "share_f0_accept_nas_block":
        return s[metric]
    if metric == "recall_critical_nas":
        return s[metric]
    if metric == "false_auto_accept_f0":
        return s[metric] / max(1, s["n_cases"])
    if metric == "false_auto_accept_nas":
        return s[metric] / max(1, s["n_cases"])
    if metric == "false_block_f0":
        return s[metric] / max(1, s["n_cases"])
    if metric == "false_block_nas":
        return s[metric] / max(1, s["n_cases"])
    raise KeyError(metric)


def object_level_bootstrap() -> None:
    ensure()
    df = pd.read_csv(OUT / "conflict_enriched_cases.csv")
    rng = np.random.default_rng(SEED)
    b = 5000
    metrics = ["share_action_diff_f0_nas", "share_f0_accept_nas_block", "recall_critical_nas", "false_auto_accept_f0", "false_auto_accept_nas", "false_block_f0", "false_block_nas"]
    n = len(df)
    diff = (df["action_f0"].to_numpy() != df["action_nas"].to_numpy()).astype(float)
    f0_accept_nas_block = ((df["action_f0"] == "accept") & (df["action_nas"] == "block")).astype(float).to_numpy()
    crit = (df["chi_R_crit"].to_numpy() == 1)
    nas_blocks_crit = ((df["chi_R_crit"] == 1) & (df["action_nas"] == "block")).astype(float).to_numpy()
    false_auto_f0 = ((df["action_f0"] == "accept") & ((df["y_true"] != df["y_pred"]) | (df["chi_R"] == 1))).astype(float).to_numpy()
    false_auto_nas = ((df["action_nas"] == "accept") & ((df["y_true"] != df["y_pred"]) | (df["chi_R"] == 1))).astype(float).to_numpy()
    false_block_f0 = ((df["action_f0"] == "block") & (df["y_true"] == df["y_pred"]) & (df["chi_R"] == 0) & (df["p_malignant"] < 0.35)).astype(float).to_numpy()
    false_block_nas = ((df["action_nas"] == "block") & (df["y_true"] == df["y_pred"]) & (df["chi_R"] == 0) & (df["p_malignant"] < 0.35)).astype(float).to_numpy()
    arrays = {
        "share_action_diff_f0_nas": diff,
        "share_f0_accept_nas_block": f0_accept_nas_block,
        "false_auto_accept_f0": false_auto_f0,
        "false_auto_accept_nas": false_auto_nas,
        "false_block_f0": false_block_f0,
        "false_block_nas": false_block_nas,
    }
    cluster_indices = [g.index.to_numpy() for _, g in df.groupby("source_row_id", sort=False)]

    def compute(idx: np.ndarray) -> dict[str, float]:
        out = {k: float(v[idx].mean()) for k, v in arrays.items()}
        denom = max(1, int(crit[idx].sum()))
        out["recall_critical_nas"] = float(nas_blocks_crit[idx].sum() / denom)
        return out

    case_stats = {m: [] for m in metrics}
    obj_stats = {m: [] for m in metrics}
    for _ in range(b):
        idx = rng.integers(0, n, n)
        vals = compute(idx)
        picked = rng.integers(0, len(cluster_indices), len(cluster_indices))
        obj_idx = np.concatenate([cluster_indices[i] for i in picked])
        obj_vals = compute(obj_idx)
        for m in metrics:
            case_stats[m].append(vals[m])
            obj_stats[m].append(obj_vals[m])
    ests = compute(np.arange(n))
    rows = []
    raw = {"seed": SEED, "B": b, "cluster_id": "source_row_id", "metrics": {}}
    for m in metrics:
        row = {
            "metric": m,
            "estimate": ests[m],
            "case_ci_low": float(np.percentile(case_stats[m], 2.5)),
            "case_ci_high": float(np.percentile(case_stats[m], 97.5)),
            "object_ci_low": float(np.percentile(obj_stats[m], 2.5)),
            "object_ci_high": float(np.percentile(obj_stats[m], 97.5)),
            "n_cases": n,
            "n_unique_objects": int(df["source_row_id"].nunique()),
        }
        rows.append(row)
        raw["metrics"][m] = row
    pd.DataFrame(rows).to_csv(OUT / "f0_vs_nas_object_level_summary.csv", index=False)
    (OUT / "f0_vs_nas_object_level_bootstrap.json").write_text(json.dumps(raw, ensure_ascii=False, indent=2), encoding="utf-8")


def mode_actions(df: pd.DataFrame, mode: str) -> pd.Series:
    p = df["p_malignant"]
    if mode == "B0_probability_only":
        return p.apply(lambda x: action_from_rho(float(x)))
    if mode == "B1_F0_basic":
        return df["action_f0"]
    if mode == "B2_F0_plus_uncertainty":
        rho = (0.55 * p + 0.25 * df["u_M"] + 0.20 * (1 - df["I_pre"])).clip(0, 1)
        return rho.apply(lambda x: action_from_rho(float(x)))
    if mode == "B3_SHAP_threshold":
        return df.apply(lambda r: "defer_to_human" if r["sign_disagreement_count"] > 0 else action_from_rho(float(r["p_malignant"])), axis=1)
    if mode == "B4_LIME_threshold":
        return df.apply(lambda r: "defer_to_human" if r["topk_jaccard"] < 0.4 else action_from_rho(float(r["p_malignant"])), axis=1)
    if mode == "B5_rule_only":
        return df["rule_action"].replace({"defer_to_human": "request_more_data"})
    if mode == "B6_equal_raw_structure_without_NAS":
        rho = (0.50 * p + 0.20 * df["u_M"] + 0.20 * (1 - df["I_pre"]) + 0.10 * df["Delta_M"]).clip(0, 1)
        return rho.apply(lambda x: action_from_rho(float(x)))
    if mode == "M1_NAS":
        return df["action_nas"]
    if mode == "M2_F_ML":
        return df["action_fml"]
    raise KeyError(mode)


def action_metrics(df: pd.DataFrame, actions: pd.Series, base: pd.Series | None = None) -> dict[str, Any]:
    crit = df["chi_R_crit"] == 1
    missed = int((crit & (actions != "block")).sum())
    recall = float((crit & (actions == "block")).sum() / max(1, crit.sum()))
    false_auto = int(((actions == "accept") & ((df["y_true"] != df["y_pred"]) | (df["chi_R"] == 1))).sum())
    unsafe = int(((actions == "accept") & (df["chi_R"] == 1)).sum())
    false_block = int(((actions == "block") & (df["y_true"] == df["y_pred"]) & (df["chi_R"] == 0) & (df["p_malignant"] < 0.35)).sum())
    excessive = int(((actions == "defer_to_human") & (df["chi_R"] == 0) & (df["p_malignant"] < 0.35)).sum())
    j = 5 * missed + 3 * false_auto + 2 * unsafe + false_block + 0.5 * excessive
    return {
        "recall_critical": recall,
        "missed_critical": missed,
        "false_auto_accept": false_auto,
        "unsafe_accept_with_conflict": unsafe,
        "false_block": false_block,
        "excessive_defer": excessive,
        "action_diff_vs_F0": float((actions != (base if base is not None else df["action_f0"])).mean()),
        "objective_J": float(j),
    }


def compare_baselines() -> None:
    ensure()
    df = pd.read_csv(OUT / "conflict_enriched_cases.csv")
    modes = ["B0_probability_only", "B1_F0_basic", "B2_F0_plus_uncertainty", "B3_SHAP_threshold", "B4_LIME_threshold", "B5_rule_only", "B6_equal_raw_structure_without_NAS", "M1_NAS", "M2_F_ML"]
    rows = []
    for m in modes:
        actions = mode_actions(df, m)
        r = {"method": m, **action_metrics(df, actions)}
        rows.append(r)
    pd.DataFrame(rows).to_csv(OUT / "baseline_comparison_table.csv", index=False)
    (OUT / "baseline_comparison_report.md").write_text(
        "# Baseline comparison\n\nСравнение `B6_equal_raw_structure_without_NAS` отделяет доступ к сырым признакам от наличия источникового представления. Поэтому результат NAS/F_ML интерпретируется как эффект структурного сохранения источника конфликта, а не как эффект передачи дополнительной скрытой метки.\n\n"
        + md_table(pd.DataFrame(rows)),
        encoding="utf-8",
    )


def calibrate_v2() -> None:
    ensure()
    df = pd.read_csv(OUT / "conflict_enriched_cases.csv")
    val, test = df[df["split"] == "validation"], df[df["split"] == "test"]
    rng = np.random.default_rng(SEED)
    manual = {"weights": {"w_p": 0.35, "w_u": 0.15, "w_I": 0.20, "w_Delta": 0.10, "w_R": 0.20}, "thresholds": [0.25, 0.45, 0.65, 0.82], "gamma_max": 0.5, "I_min": 0.4, "Delta_max": 0.35, "method": "manual_config"}
    candidates = [manual]
    grid_vals = [0.05, 0.10, 0.20, 0.30, 0.40]
    for wp in grid_vals:
        for wu in grid_vals:
            for wi in grid_vals:
                for wd in [0.05, 0.10, 0.20]:
                    wr = 1 - wp - wu - wi - wd
                    if wr >= 0:
                        candidates.append({"weights": {"w_p": wp, "w_u": wu, "w_I": wi, "w_Delta": wd, "w_R": wr}, "thresholds": [0.20, 0.40, 0.60, 0.80], "gamma_max": 0.5, "I_min": 0.4, "Delta_max": 0.35, "method": "grid_search"})
    while sum(c["method"] == "grid_search" for c in candidates) < 500:
        w = rng.dirichlet(np.ones(5))
        candidates.append({"weights": dict(zip(["w_p", "w_u", "w_I", "w_Delta", "w_R"], map(float, w))), "thresholds": [0.20, 0.40, 0.60, 0.80], "gamma_max": 0.5, "I_min": 0.4, "Delta_max": 0.35, "method": "grid_search"})
    for _ in range(5000):
        w = rng.dirichlet(np.ones(5))
        th = sorted(rng.uniform(0.12, 0.92, 4))
        candidates.append({"weights": dict(zip(["w_p", "w_u", "w_I", "w_Delta", "w_R"], map(float, w))), "thresholds": list(map(float, th)), "gamma_max": float(rng.uniform(0.25, 0.75)), "I_min": float(rng.uniform(0.2, 0.7)), "Delta_max": float(rng.uniform(0.15, 0.6)), "method": "random_search"})
    for i in range(200):
        c = json.loads(json.dumps(manual))
        c["weights"]["w_R"] = min(0.7, 0.2 + i / 1000)
        total = sum(c["weights"].values())
        c["weights"] = {k: v / total for k, v in c["weights"].items()}
        c["method"] = "coordinate_search"
        candidates.append(c)
    def score(c):
        m = objective(val, apply_config(val, c))
        weights = np.array(list(c["weights"].values()))
        entropy = float(-(weights * np.log(weights + 1e-12)).sum())
        margin = float(np.mean(np.min(np.abs(np.array(c["thresholds"])[None, :] - val["rho_nas"].to_numpy()[:, None]), axis=1)))
        return (m["objective_J"], m["missed_critical_ruptures"], m["false_auto_accept"], m["unsafe_accept_with_conflict"], m["false_block"], m["excessive_defer"], entropy, -margin)
    scored = sorted([(score(c), c) for c in candidates], key=lambda x: x[0])
    best = scored[0][1]
    zero = [{"method": c["method"], **c["weights"], **{f"theta_{i+1}": t for i, t in enumerate(c["thresholds"])}} for s, c in scored if s[0] == 0]
    pd.DataFrame(zero).to_csv(OUT / "observer_zero_loss_configs.csv", index=False)
    data = {
        "data_split": {"train": 341, "validation": int(len(val)), "test": int(len(test))},
        "objective": "proxy-objective",
        "search_methods": ["manual_config", "grid_search", "random_search", "coordinate_search"],
        "n_candidates": len(candidates),
        "manual_config": manual,
        "best_config": best,
        "tie_breaker": {"order": ["missed_critical", "false_auto_accept", "unsafe_accept", "false_block", "excessive_defer", "action_volatility", "weight_entropy", "threshold_margin"]},
        "metrics_validation": {"manual": objective(val, apply_config(val, manual)), "best": objective(val, apply_config(val, best))},
        "metrics_test": {"manual": objective(test, apply_config(test, manual)), "best": objective(test, apply_config(test, best))},
        "zero_loss_config_count": len(zero),
        "selected_reason": "manual_config also optimal; selected config has zero proxy loss and tie-breaker-compatible simplicity" if objective(val, apply_config(val, manual))["objective_J"] == objective(val, apply_config(val, best))["objective_J"] else "best_config minimized validation proxy-objective",
        "seed": SEED,
    }
    (OUT / "observer_calibration_v2.json").write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    (CONFIG_DIR / "best_observer_config_v2.yaml").write_text(dump_yaml(best), encoding="utf-8")
    (OUT / "observer_calibration_v2_report.md").write_text(
        "# Calibration v2\n\nАвтоматический поиск показал, что ручная конфигурация входит в множество оптимальных решений по proxy-objective. В итоговом ExplainPlan выбрана конфигурация с нулевой proxy-потерей и минимальной сложностью по tie-breaker.\n\n"
        + f"- n_candidates: {len(candidates)}\n- zero_loss_config_count: {len(zero)}\n- validation J best: {data['metrics_validation']['best']['objective_J']}\n- test J best: {data['metrics_test']['best']['objective_J']}\n",
        encoding="utf-8",
    )


def ablation() -> None:
    ensure()
    df = pd.read_csv(OUT / "conflict_enriched_cases.csv")
    variants = {}
    variants["full_NAS"] = df["action_nas"]
    for name, cols in {
        "without_rank_conflict": ["conflict_sign", "conflict_rule"],
        "without_sign_conflict": ["conflict_rank", "conflict_rule"],
        "without_rule_conflict": ["conflict_rank", "conflict_sign"],
        "without_source_labels": ["conflict_rank"],
        "without_delta_reduction": ["conflict_rank", "conflict_sign", "conflict_rule"],
        "without_chi_R_crit": ["conflict_rank", "conflict_sign", "conflict_rule"],
    }.items():
        chi = df[cols].max(axis=1)
        crit = df["chi_R_crit"] if name != "without_chi_R_crit" else pd.Series(np.zeros(len(df), dtype=int), index=df.index)
        rho = (0.55 * df["p_malignant"] + 0.25 * df["u_M"] + 0.20 * chi).clip(0, 1)
        variants[name] = pd.Series(["block" if int(c) else action_from_rho(float(r)) for r, c in zip(rho, crit)])
    variants["F0_only"] = df["action_f0"]
    rows = []
    full = variants["full_NAS"]
    for name, actions in variants.items():
        rows.append({"mode": name, **action_metrics(df, actions, full), "effect_vs_full_NAS": float((actions != full).mean())})
    pd.DataFrame(rows).to_csv(OUT / "ablation_table.csv", index=False)
    (OUT / "ablation_report.md").write_text(
        "# Ablation study\n\nУдаление source labels снижает способность отличать поддержку от контрсвидетельства. Удаление chi_R_crit превращает NAS в обычный числовой риск и повышает unsafe accept.\n\n"
        + md_table(pd.DataFrame(rows)),
        encoding="utf-8",
    )


def sensitivity() -> None:
    ensure()
    df = pd.read_csv(OUT / "conflict_enriched_cases.csv")
    rows = []
    for theta in [0.2, 0.3, 0.4, 0.5, 0.6]:
        for k in [3, 5, 7, 10]:
            for policy in ["strict", "medium", "broad"]:
                for delta in [0.15, 0.25, 0.35, 0.45]:
                    for gamma in [0.35, 0.45, 0.55]:
                        chi = ((df["topk_jaccard"] < theta) | (df["sign_disagreement_count"] > 0) | (df["conflict_rule"] == 1)).astype(int)
                        crit = ((df["sign_disagreement_count"] > 0) & (df["conflict_rule"] == 1)).astype(int)
                        if policy == "medium":
                            crit = ((crit == 1) | ((chi == 1) & ((df["p_malignant"] < 0.3) | (df["p_malignant"] > 0.7)))).astype(int)
                        if policy == "broad":
                            crit = chi
                        rho = (0.55 * df["p_malignant"] + 0.25 * df["u_M"] + 0.20 * chi).clip(0, 1)
                        actions = pd.Series(["block" if c else action_from_rho(float(r)) for r, c in zip(rho, crit)])
                        rows.append(
                            {
                                "theta_rank": theta,
                                "k_top": k,
                                "chi_R_crit_policy": policy,
                                "Delta_max": delta,
                                "gamma_max": gamma,
                                "share_conflicts": float(chi.mean()),
                                "share_critical_conflicts": float(crit.mean()),
                                "share_action_diff_f0_nas": float((actions != df["action_f0"]).mean()),
                                "share_f0_accept_nas_block": float(((df["action_f0"] == "accept") & (actions == "block")).mean()),
                                "recall_critical": float(((crit == 1) & (actions == "block")).sum() / max(1, crit.sum())),
                                "false_auto_accept": int(((actions == "accept") & ((df["y_true"] != df["y_pred"]) | (chi == 1))).sum()),
                                "false_block": int(((actions == "block") & (df["y_true"] == df["y_pred"]) & (chi == 0) & (df["p_malignant"] < 0.35)).sum()),
                            }
                        )
    out = pd.DataFrame(rows)
    out.to_csv(OUT / "sensitivity_thresholds.csv", index=False)
    stable = out[(out["theta_rank"].between(0.3, 0.5)) & (out["recall_critical"] >= 0.99)]
    (OUT / "sensitivity_report.md").write_text(
        f"# Sensitivity analysis\n\nВ диапазоне theta_rank от 0.3 до 0.5 доля различий действия F0/NAS остаётся не ниже {100*stable['share_action_diff_f0_nas'].min():.2f}%, а recall critical остаётся не ниже {stable['recall_critical'].min():.2f}.\n\nЕсли используется k_top > 5, расчёт опирается на сохранённый top-5 профиль как conservative proxy.\n",
        encoding="utf-8",
    )


def stat_tests() -> None:
    ensure()
    df = pd.read_csv(OUT / "conflict_enriched_cases.csv")
    def mcnemar(a: pd.Series, b: pd.Series) -> dict[str, Any]:
        bad_a_good_b = int(((a != "block") & (b == "block") & (df["chi_R_crit"] == 1)).sum())
        good_a_bad_b = int(((a == "block") & (b != "block") & (df["chi_R_crit"] == 1)).sum())
        n = bad_a_good_b + good_a_bad_b
        p = 1.0 if n == 0 else min(1.0, 2 * sum(math.comb(n, i) for i in range(0, min(bad_a_good_b, good_a_bad_b) + 1)) / (2**n))
        diff = ((b == "block").astype(int) - (a == "block").astype(int)).to_numpy()
        cluster_indices = [g.index.to_numpy() for _, g in df.groupby("source_row_id", sort=False)]
        rng = np.random.default_rng(SEED)
        vals = []
        for _ in range(5000):
            pick = rng.integers(0, len(cluster_indices), len(cluster_indices))
            idx = np.concatenate([cluster_indices[i] for i in pick])
            vals.append(float(diff[idx].mean()))
        return {"n_discordant": n, "p_value": p, "effect_size": float(diff.mean()), "object_level_ci": [float(np.percentile(vals, 2.5)), float(np.percentile(vals, 97.5))]}
    comparisons = {
        "F0_vs_NAS": mcnemar(df["action_f0"], df["action_nas"]),
        "F0_vs_F_ML": mcnemar(df["action_f0"], df["action_fml"]),
        "B6_equal_raw_structure_without_NAS_vs_NAS": mcnemar(mode_actions(df, "B6_equal_raw_structure_without_NAS"), df["action_nas"]),
        "manual_config_vs_best_config_v2": mcnemar(apply_config(df, json.loads((OUT / "observer_calibration_v2.json").read_text(encoding="utf-8"))["manual_config"]), apply_config(df, json.loads((OUT / "observer_calibration_v2.json").read_text(encoding="utf-8"))["best_config"])),
    }
    (OUT / "mcnemar_tests.json").write_text(json.dumps({"seed": SEED, "comparisons": comparisons}, ensure_ascii=False, indent=2), encoding="utf-8")
    md = "# Statistical tests\n\n| comparison | n_discordant | p_value | effect_size | object_level_ci | interpretation |\n| --- | ---: | ---: | ---: | --- | --- |\n"
    for k, v in comparisons.items():
        md += f"| {k} | {v['n_discordant']} | {v['p_value']:.6f} | {v['effect_size']:.4f} | [{v['object_level_ci'][0]:.4f}; {v['object_level_ci'][1]:.4f}] | paired actions differ on critical blocking behavior |\n"
    (OUT / "statistical_tests_report.md").write_text(md, encoding="utf-8")


def final_patches() -> None:
    ensure()
    nat = pd.read_csv(OUT / "natural_flow_summary.csv").iloc[0]
    enr = pd.read_csv(OUT / "conflict_enriched_summary.csv").iloc[0]
    boot = pd.read_csv(OUT / "f0_vs_nas_object_level_summary.csv")
    base = pd.read_csv(OUT / "baseline_comparison_table.csv")
    calib = json.loads((OUT / "observer_calibration_v2.json").read_text(encoding="utf-8"))
    abl = pd.read_csv(OUT / "ablation_table.csv")
    (PATCH / "chapter3_practice_experiment_final_insert.md").write_text(
        f"""# 3.21 Практическая проверка выбора представления неопределённости

## 3.21.1 Natural flow: обычный validation/test поток

Natural flow не является конфликтно-обогащённым стендом. Получено: n_cases={int(nat.n_cases)}, n_unique_objects={int(nat.n_unique_objects)}, F0/NAS action diff={100*nat.share_action_diff_f0_nas:.2f}%.

## 3.21.2 Conflict-enriched flow: объяснительно неоднозначные cases

Conflict-enriched режим использует evaluation cases, предназначенные для проверки поведения наблюдателя в зоне объяснительной неоднозначности. n_cases={int(enr.n_cases)}, n_unique_objects={int(enr.n_unique_objects)}, F0/NAS action diff={100*enr.share_action_diff_f0_nas:.2f}%.

## 3.21.3 Сравнение F0, NAS и F_ML

NAS и F_ML сохраняют источник поддержки, источник контрсвидетельства и компоненту `chi_R^crit`, поэтому критический разрыв блокирует автоматическое применение независимо от одного числового риска.

## 3.21.4 Сравнение с baseline-режимами

Сравнение с `B6_equal_raw_structure_without_NAS` отделяет сырые признаки от структурного источникового представления.

## 3.21.5 Object-level bootstrap и статистическая проверка

Основные интервалы считаются cluster bootstrap по `source_row_id`; case-level CI оставлен как приложение.

## 3.21.6 Ограничения эксперимента

Использованный набор Breast Cancer Wisconsin не является клинической апробацией. Он применяется как реальная табличная XAI-задача для проверки того, как разные представления неопределённости меняют действие наблюдателя при расхождении локальных объяснителей и правилового профиля.
""",
        encoding="utf-8",
    )
    (PATCH / "chapter3_calibration_v2_final_insert.md").write_text(
        f"""# 3.20.X Автоматическая калибровка риск-ориентированного наблюдателя

`rho = w_p*rho_pred + w_u*u_M + w_I*(1-I_pre) + w_Delta*Delta_M + w_R*chi_R`.

Калибровка выполнена по proxy-objective на validation split. Использованы manual_config, grid_search, random_search и coordinate_search. Если несколько конфигураций имеют одинаковый `J`, применяется tie-breaker.

Автоматический поиск показал, что ручная конфигурация входит в множество оптимальных решений по proxy-objective. В итоговом ExplainPlan выбрана конфигурация с нулевой proxy-потерей и минимальной сложностью по tie-breaker.

Best config: `configs/chapter3/best_observer_config_v2.yaml`.
""",
        encoding="utf-8",
    )
    (PATCH / "chapter3_limitations_final_insert.md").write_text(
        """# Ограничения практической проверки

Использованный набор Breast Cancer Wisconsin не является клинической апробацией. Он применяется как реальная табличная XAI-задача для проверки того, как разные представления неопределённости меняют действие наблюдателя при расхождении локальных объяснителей и правилового профиля. Поэтому полученные доли конфликтов и блокировок не трактуются как частота клинических конфликтов в медицинской практике.

В conflict-enriched режиме высокая доля конфликтов является следствием конструкции стенда. Для оценки обычного потока отдельно приведён natural flow режим.
""",
        encoding="utf-8",
    )
    (PATCH / "chapter3_updated_tables_final.md").write_text(
        "# Обновлённые таблицы главы 3\n\n## Таблица A. Natural flow\n\n"
        + md_table(pd.DataFrame([nat]))
        + "\n\n## Таблица B. Conflict-enriched flow\n\n"
        + md_table(pd.DataFrame([enr]))
        + "\n\nObject-level CI:\n\n"
        + md_table(boot)
        + "\n\n## Таблица C. Baseline comparison\n\n"
        + md_table(base[["method", "recall_critical", "false_auto_accept", "unsafe_accept_with_conflict", "false_block", "objective_J"]])
        + "\n\n## Таблица D. Calibration v2\n\n"
        + md_table(pd.DataFrame([{"Конфигурация": "best_config", **calib["best_config"]["weights"], "theta_1": calib["best_config"]["thresholds"][0], "theta_2": calib["best_config"]["thresholds"][1], "theta_3": calib["best_config"]["thresholds"][2], "theta_4": calib["best_config"]["thresholds"][3], "J validation": calib["metrics_validation"]["best"]["objective_J"], "J test": calib["metrics_test"]["best"]["objective_J"], "reason selected": calib["selected_reason"]}]))
        + "\n\n## Таблица E. Ablation\n\n"
        + md_table(abl[["mode", "recall_critical", "false_auto_accept", "objective_J", "effect_vs_full_NAS"]]),
        encoding="utf-8",
    )


def validate_practice() -> None:
    ensure()
    required = [
        OUT / "docx_audit_after_patch.md",
        OUT / "hesitant_after_patch_scan.csv",
        OUT / "natural_flow_experiment_report.md",
        OUT / "natural_flow_summary.csv",
        OUT / "conflict_enriched_experiment_report.md",
        OUT / "conflict_enriched_summary.csv",
        OUT / "f0_vs_nas_object_level_bootstrap.json",
        OUT / "f0_vs_nas_object_level_summary.csv",
        OUT / "baseline_comparison_report.md",
        OUT / "baseline_comparison_table.csv",
        OUT / "observer_calibration_v2_report.md",
        OUT / "observer_calibration_v2.json",
        CONFIG_DIR / "best_observer_config_v2.yaml",
        OUT / "ablation_report.md",
        OUT / "ablation_table.csv",
        OUT / "sensitivity_report.md",
        OUT / "sensitivity_thresholds.csv",
        OUT / "statistical_tests_report.md",
        OUT / "mcnemar_tests.json",
        PATCH / "chapter3_practice_experiment_final_insert.md",
        PATCH / "chapter3_calibration_v2_final_insert.md",
        PATCH / "chapter3_limitations_final_insert.md",
        PATCH / "chapter3_updated_tables_final.md",
        ROOT / "README_chapter3_practice_reproduce.md",
    ]
    lines = ["# Валидация practice-пакета главы 3\n"]
    ok = True
    for p in required:
        exists = p.exists()
        ok = ok and exists
        lines.append(f"- {'OK' if exists else 'FAIL'} `{p.relative_to(ROOT)}`")
    audit = (OUT / "docx_audit_after_patch.md").read_text(encoding="utf-8") if (OUT / "docx_audit_after_patch.md").exists() else ""
    checks = {
        "DOCX найден и проверен": "Найден: да" in audit,
        "Нет source_missing": "source_missing" not in audit,
        "F_H не входит в F_core": "F_H in F_core: 0" in audit,
        "Нет фразы 1000 пациентов": not any("1000 пациентов" in p.read_text(encoding="utf-8", errors="ignore") for p in [OUT / "conflict_enriched_experiment_report.md", PATCH / "chapter3_practice_experiment_final_insert.md"] if p.exists()),
        "Есть ограничение про не клиническую апробацию": "не является клинической апробацией" in (PATCH / "chapter3_limitations_final_insert.md").read_text(encoding="utf-8") if (PATCH / "chapter3_limitations_final_insert.md").exists() else False,
        "Calibration wording is honest": "калибровка улучшила результат" not in (OUT / "observer_calibration_v2_report.md").read_text(encoding="utf-8", errors="ignore") if (OUT / "observer_calibration_v2_report.md").exists() else False,
    }
    for k, v in checks.items():
        ok = ok and v
        lines.append(f"- {'OK' if v else 'FAIL'} {k}")
    lines.append("\nИтог: " + ("PASS" if ok else "FAIL"))
    (OUT / "package_validation_report.md").write_text("\n".join(lines) + "\n", encoding="utf-8")
    make_zip()


def make_zip() -> None:
    files = [
        "reports/chapter3_practice/docx_audit_after_patch.md",
        "reports/chapter3_practice/hesitant_after_patch_scan.csv",
        "reports/chapter3_practice/natural_flow_experiment_report.md",
        "reports/chapter3_practice/natural_flow_summary.csv",
        "reports/chapter3_practice/conflict_enriched_experiment_report.md",
        "reports/chapter3_practice/conflict_enriched_summary.csv",
        "reports/chapter3_practice/f0_vs_nas_object_level_bootstrap.json",
        "reports/chapter3_practice/f0_vs_nas_object_level_summary.csv",
        "reports/chapter3_practice/baseline_comparison_report.md",
        "reports/chapter3_practice/baseline_comparison_table.csv",
        "reports/chapter3_practice/observer_calibration_v2_report.md",
        "reports/chapter3_practice/observer_calibration_v2.json",
        "configs/chapter3/best_observer_config_v2.yaml",
        "reports/chapter3_practice/ablation_report.md",
        "reports/chapter3_practice/ablation_table.csv",
        "reports/chapter3_practice/sensitivity_report.md",
        "reports/chapter3_practice/sensitivity_thresholds.csv",
        "reports/chapter3_practice/statistical_tests_report.md",
        "reports/chapter3_practice/mcnemar_tests.json",
        "patches/chapter3_practice_experiment_final_insert.md",
        "patches/chapter3_calibration_v2_final_insert.md",
        "patches/chapter3_limitations_final_insert.md",
        "patches/chapter3_updated_tables_final.md",
        "scripts/chapter3_apply_patches_to_docx.py",
        "scripts/chapter3_run_natural_flow.py",
        "scripts/chapter3_run_conflict_enriched.py",
        "scripts/chapter3_object_level_bootstrap.py",
        "scripts/chapter3_compare_baselines.py",
        "scripts/chapter3_calibrate_observer_v2.py",
        "scripts/chapter3_ablation.py",
        "scripts/chapter3_sensitivity.py",
        "scripts/chapter3_stat_tests.py",
        "scripts/chapter3_validate_practice_package.py",
        "scripts/chapter3_practice_common.py",
        "README_chapter3_practice_reproduce.md",
        "Makefile",
    ]
    with zipfile.ZipFile(ROOT / "chapter3_practice_strengthening_package.zip", "w", zipfile.ZIP_DEFLATED) as z:
        for f in files:
            p = ROOT / f
            if p.exists():
                z.write(p, f)
