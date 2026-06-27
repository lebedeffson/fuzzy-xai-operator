#!/usr/bin/env python3
"""Patch chapter 2 DOCX according to package 2 requirements."""

from __future__ import annotations

import argparse
import re
import shutil
from pathlib import Path

from docx import Document
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_INPUTS = [
    ROOT / "docs/chapter2/glava2_chapter2_package1_fixed.docx",
    Path("/mnt/data/glava2_chapter2_package1_fixed.docx"),
]
OUTPUT = ROOT / "glava2_chapter2_package2_full_fixed.docx"
SECTION = ROOT / "patches/section_2_18_2.md"
FIG_DIR = ROOT / "figures/chapter2"
PNG_DIR = FIG_DIR / "png_preview"

FIGURES = [
    {
        "num": "2.1",
        "file": "fig_2_1_Ek_structure",
        "caption": "Рисунок 2.1. Структура нечёткого объяснительного объекта E_k.",
        "anchor": ("2.5", "E_k"),
        "ref": "Структура объекта E_k показана на рисунке 2.1.",
    },
    {
        "num": "2.2",
        "file": "fig_2_2_Tij_composition",
        "caption": "Рисунок 2.2. Согласование T_ij и частичная композиция объяснительных объектов.",
        "anchor": ("2.7", None),
        "ref": "Логика согласования T_ij и выбора между E_ij и D_ij показана на рисунке 2.2.",
    },
    {
        "num": "2.3",
        "file": "fig_2_3_dE_metric",
        "caption": "Рисунок 2.3. Компонентная структура метрической нормальной формы d_E.",
        "anchor": ("2.8.1", "d_E"),
        "ref": "Компонентная структура функционала d_E показана на рисунке 2.3.",
    },
    {
        "num": "2.4",
        "file": "fig_2_4_GExpl_graph",
        "caption": "Рисунок 2.4. Граф сертифицированных объяснительных переходов G_Expl.",
        "anchor": ("2.7.2", None),
        "ref": "Маршрут сертифицированных переходов G_Expl показан на рисунке 2.4.",
    },
    {
        "num": "2.5",
        "file": "fig_2_5_sample113_flow",
        "caption": "Рисунок 2.5. Сквозной численный маршрут для объекта sample_113.",
        "anchor": ("sample_113", None),
        "ref": "Сквозной маршрут для sample_113 показан на рисунке 2.5.",
    },
]

REPLACEMENTS = {
    "для объяснимый ИИ": "для XAI",
    "методы объяснимый ИИ": "методы XAI",
    "в пространстве объяснительного ИИ": "в пространстве XAI",
    "объяснимого ИИ": "XAI",
    "объяснимому ИИ": "XAI",
    "объяснимым ИИ": "XAI",
    "объяснимый ИИ": "XAI",
}


def choose_input(path: str | None) -> Path:
    candidates = [Path(path)] if path else DEFAULT_INPUTS
    for p in candidates:
        if p.exists():
            local = ROOT / "docs/chapter2/glava2_chapter2_package1_fixed.docx"
            local.parent.mkdir(parents=True, exist_ok=True)
            if p.resolve() != local.resolve():
                shutil.copy2(p, local)
            return local
    raise FileNotFoundError(
        "Не найден исходный DOCX: docs/chapter2/glava2_chapter2_package1_fixed.docx "
        "или /mnt/data/glava2_chapter2_package1_fixed.docx"
    )


def para_text(p):
    return "".join(run.text for run in p.runs) if p.runs else p.text


def replace_terms_in_paragraph(p):
    for run in p.runs:
        text = run.text
        for old, new in REPLACEMENTS.items():
            text = text.replace(old, new)
        run.text = text
    full = para_text(p)
    fixed = full
    for old, new in REPLACEMENTS.items():
        fixed = fixed.replace(old, new)
    if fixed != full:
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = fixed
        else:
            p.add_run(fixed)


def replace_terms(doc):
    for p in doc.paragraphs:
        replace_terms_in_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_terms_in_paragraph(p)


def remove_paragraph(p):
    el = p._element
    el.getparent().remove(el)


def has_drawing(p):
    return bool(p._element.xpath(".//w:drawing|.//w:pict"))


def remove_old_figures(doc):
    paras = list(doc.paragraphs)
    remove_ids = set()
    for i, p in enumerate(paras):
        if re.match(r"^\s*Рисунок\s+2\.[1-5]\b", para_text(p)):
            remove_ids.add(i)
            j = i - 1
            while j >= 0 and not para_text(paras[j]).strip():
                remove_ids.add(j)
                j -= 1
            if j >= 0 and has_drawing(paras[j]):
                remove_ids.add(j)
    for i in sorted(remove_ids, reverse=True):
        remove_paragraph(paras[i])


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = paragraph.insert_paragraph_before(text, style=style)
    paragraph._element.addnext(new_p._element)
    return new_p


def insert_paragraph_before(paragraph, text="", style=None):
    return paragraph.insert_paragraph_before(text, style=style)


def find_heading(doc, prefix):
    pat = re.compile(rf"^\s*{re.escape(prefix)}(\s|$)")
    for i, p in enumerate(doc.paragraphs):
        if pat.search(para_text(p)):
            return i
    return None


def find_anchor(doc, spec):
    key, sub = spec
    if key == "sample_113":
        for i, p in enumerate(doc.paragraphs):
            if "sample_113" in para_text(p):
                return i
        return len(doc.paragraphs) - 1
    h = find_heading(doc, key)
    if h is None:
        return len(doc.paragraphs) - 1
    if not sub:
        return h
    for i in range(h + 1, min(len(doc.paragraphs), h + 25)):
        t = para_text(doc.paragraphs[i])
        if sub in t:
            return i
    return h


def add_figures(doc):
    for fig in FIGURES:
        img = PNG_DIR / f"{fig['file']}.png"
        if not img.exists():
            raise FileNotFoundError(f"Нет PNG-превью для вставки: {img}")
        idx = find_anchor(doc, fig["anchor"])
        anchor = doc.paragraphs[idx]
        refp = insert_paragraph_after(anchor, fig["ref"])
        picp = insert_paragraph_after(refp)
        picp.alignment = 1
        picp.add_run().add_picture(str(img), width=Inches(6.3))
        cap = insert_paragraph_after(picp, fig["caption"])
        cap.alignment = 1


def section_paragraphs():
    raw = SECTION.read_text(encoding="utf-8").splitlines()
    out = []
    for line in raw:
        line = line.strip()
        if not line:
            continue
        out.append(line[2:] if line.startswith("# ") else line)
    return out


def add_section_2_18_2(doc):
    if any("2.18.2 Адаптация оператора к разным XAI-методам" in para_text(p) for p in doc.paragraphs):
        return
    start = find_heading(doc, "2.18.1")
    if start is None:
        target = len(doc.paragraphs) - 1
        p = doc.paragraphs[target]
        for text in section_paragraphs():
            p = insert_paragraph_after(p, text)
        return
    insert_before = None
    for i in range(start + 1, len(doc.paragraphs)):
        t = para_text(doc.paragraphs[i]).strip()
        if re.match(r"^2\.(18\.2|19|20|\d{2,})(\s|$)", t):
            insert_before = i
            break
    if insert_before is None:
        p = doc.paragraphs[-1]
        for text in section_paragraphs():
            p = insert_paragraph_after(p, text)
    else:
        first = True
        for text in section_paragraphs():
            if first:
                p = insert_paragraph_before(doc.paragraphs[insert_before], text)
                first = False
            else:
                p = insert_paragraph_after(p, text)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", help="Source DOCX")
    parser.add_argument("--output", default=str(OUTPUT), help="Output DOCX")
    args = parser.parse_args()
    src = choose_input(args.input)
    doc = Document(str(src))
    remove_old_figures(doc)
    replace_terms(doc)
    add_section_2_18_2(doc)
    add_figures(doc)
    out = Path(args.output)
    doc.save(str(out))
    print(f"Saved {out}")


if __name__ == "__main__":
    main()
